#!/usr/bin/env python3
"""Export every proposal in the vault index to a real .docx.

    python backend/build_vault_index.py        # scan
    python backend/tools/proposals_to_docx.py  # export

Writes deliverables/proposals/docx/<YYYY-MM-DD_Proposal_Slug>.docx plus a
manifest the dashboard reads, so each proposal gets an "OPEN .DOCX" button
served straight off Railway — no Google Drive required.
"""
from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    raise SystemExit("pip install python-docx")

ROOT = Path(__file__).resolve().parent.parent.parent
INDEX = ROOT / "backend" / "outputs" / "vault-index.json"
OUTDIR = ROOT / "deliverables" / "proposals" / "docx"
MANIFEST = ROOT / "backend" / "outputs" / "docx-manifest.json"

GOLD = RGBColor(0x9A, 0x7D, 0x1E)
INK = RGBColor(0x1A, 0x1A, 0x1A)


def slug(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "-", s).strip("-")[:70]


def _style(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10.5)
    n.font.color.rgb = INK


def _para(doc: Document, text: str, size=10.5, bold=False, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def _md_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=width)
    t.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(width):
            txt = row[j] if j < len(row) else ""
            cells[j].text = txt
            for p in cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.bold = i == 0


def _render_md(doc: Document, md: str) -> None:
    lines, i = md.split("\n"), 0
    while i < len(lines):
        ln = lines[i]
        if re.match(r"^\s*\|.*\|\s*$", ln):
            rows = []
            while i < len(lines) and re.match(r"^\s*\|.*\|\s*$", lines[i]):
                cells = [c.replace("\\|", "|").strip()
                         for c in re.split(r"(?<!\\)\|", lines[i].strip().strip("|"))]
                rows.append(cells)
                i += 1
            rows = [r for r in rows if not all(re.fullmatch(r"-{2,}|", c or "") for c in r)]
            _md_table(doc, rows)
            doc.add_paragraph()
            continue
        if re.match(r"^\s*[-*]\s+", ln):
            doc.add_paragraph(re.sub(r"^\s*[-*]\s+", "", ln).replace("**", ""), style="List Bullet")
            i += 1
            continue
        if re.match(r"^#{3,4}\s+", ln):
            _para(doc, re.sub(r"^#+\s+", "", ln).replace("**", ""), size=10.5, bold=True, color=GOLD, space_after=4)
            i += 1
            continue
        if ln.strip():
            _para(doc, ln.replace("**", ""))
        i += 1


def build_one(p: dict) -> Path | None:
    doc = Document()
    _style(doc)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = h.add_run(p["title"])
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = INK

    meta = " · ".join(x for x in (p.get("sector"), p.get("market"), p.get("type")) if x)
    if meta:
        _para(doc, meta, size=9.5, color=GOLD, space_after=2)
    _para(doc, "ZYNTH · The Intelligence of Creativity · zynth.asia", size=8.5, space_after=14)

    if p.get("sections"):
        for s in p["sections"]:
            if s["level"] == 2:
                _para(doc, s["h"], size=13.5, bold=True, color=INK, space_after=4)
            else:
                _para(doc, s["h"], size=11, bold=True, color=GOLD, space_after=3)
            _render_md(doc, s["md"])
    else:
        blocks = [("The big idea", p.get("idea")),
                  ("Key deliverables", p.get("deliverables")),
                  ("KPIs", p.get("kpis")),
                  ("Investment", p.get("investment")),
                  ("The ZYNTH edge", p.get("edge"))]
        if not any(v for _, v in blocks):
            return None
        for label, val in blocks:
            if not val:
                continue
            _para(doc, label.upper(), size=10, bold=True, color=GOLD, space_after=3)
            if isinstance(val, list):
                for v in val:
                    doc.add_paragraph(v, style="List Bullet")
                doc.add_paragraph()
            else:
                _para(doc, str(val), space_after=10)

    _para(doc, "", space_after=10)
    _para(doc, "Prepared by ZYNTH. Figures are indicative and subject to vendor RFQ. "
               "Market FX applies. 50% deposit on signature.", size=8, space_after=0)

    OUTDIR.mkdir(parents=True, exist_ok=True)
    name = f"{date.today().isoformat()}_Proposal_{slug(p['title'])}.docx"
    path = OUTDIR / name
    doc.save(path)
    return path


def main() -> None:
    data = json.loads(INDEX.read_text(encoding="utf-8"))
    out, skipped = {}, 0
    for p in data["proposals"]:
        if not p.get("composed"):
            skipped += 1
            continue
        path = build_one(p)
        if path:
            out[p["title"]] = {"file": path.name,
                               "url": f"/docs/{path.name}",
                               "kb": round(path.stat().st_size / 1024, 1),
                               "full": bool(p.get("full"))}
    MANIFEST.write_text(json.dumps(out, ensure_ascii=False, indent=1), encoding="utf-8")
    print(f"wrote {len(out)} .docx to {OUTDIR.relative_to(ROOT)}")
    for t, m in out.items():
        print(f"  {'FULL' if m['full'] else '    '}  {m['kb']:6.1f} KB  {t[:56]}")
    print(f"skipped {skipped} concepts (nothing written to export yet)")


if __name__ == "__main__":
    main()
