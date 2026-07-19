"""Word document generation for client-facing ZYNTH proposals.

Renders a master proposal (title + numbered sections + tables) into a
premium-looking .docx per the ZYNTH document standard: navy/gold brand
styling, a one-line-ask callout on the cover, navy table header rows with
white text, and zebra-striped body rows. No emoji, numbered sections,
client-grade formatting. Files are written to outputs/proposals/.
"""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_OUT_DIR = Path("outputs/proposals")

# ZYNTH brand palette
_NAVY = RGBColor(0x0B, 0x1F, 0x3A)     # deep navy — headings, table header fill
_GOLD = RGBColor(0xB8, 0x8A, 0x2A)     # metallic gold — accents, rules
_CYAN = RGBColor(0x0E, 0x7C, 0x86)     # sub-heads
_INK = RGBColor(0x1A, 0x1A, 0x1A)      # body
_GREY = RGBColor(0x55, 0x5A, 0x63)     # meta / notes
_ZEBRA = "F2F1EC"                       # warm light band for alternate rows
_NAVY_HEX = "0B1F3A"


def _safe_filename(text: str) -> str:
    text = re.sub(r"[^\w\s-]", "", text).strip()
    return re.sub(r"[\s]+", "_", text)[:80] or "proposal"


def _shade(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:color"), "auto")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def _rule(doc, color_hex: str = "B88A2A") -> None:
    """A thin bottom-border paragraph used as a horizontal rule."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pbdr.append(bottom)
    pPr.append(pbdr)


def build_proposal_docx(
    title: str,
    client: str,
    market: str,
    sections: list[dict[str, str]],
    prepared_by: str = "ZYNTH — The Intelligence of Creativity",
    one_line_ask: str = "",
    estimated_value: str = "",
) -> Path:
    """Render the proposal to a branded .docx and return the file path."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ── Cover block ────────────────────────────────────────────────────────
    t = doc.add_paragraph()
    t.paragraph_format.space_after = Pt(2)
    tr = t.add_run(title)
    tr.bold = True
    tr.font.size = Pt(26)
    tr.font.color.rgb = _NAVY

    sub = doc.add_paragraph()
    sr = sub.add_run("The Intelligence of Creativity")
    sr.italic = True
    sr.font.size = Pt(11)
    sr.font.color.rgb = _GOLD
    _rule(doc)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_before = Pt(4)
    for label, value in (
        ("Prepared for", client),
        ("Market", market),
        ("Date", datetime.now().strftime("%d %B %Y")),
        ("Prepared by", prepared_by),
    ):
        r = meta.add_run(f"{label}:  ")
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = _NAVY
        v = meta.add_run(f"{value}\n")
        v.font.size = Pt(9.5)
        v.font.color.rgb = _GREY

    # One-line-ask callout — the headline the client reads first
    if one_line_ask:
        box = doc.add_paragraph()
        box.paragraph_format.space_before = Pt(6)
        box.paragraph_format.left_indent = Inches(0.15)
        pPr = box._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), "B88A2A")
        pbdr.append(left)
        pPr.append(pbdr)
        br = box.add_run(one_line_ask)
        br.bold = True
        br.font.size = Pt(11)
        br.font.color.rgb = _NAVY
    elif estimated_value:
        ev = doc.add_paragraph()
        evr = ev.add_run(f"Estimated investment: {estimated_value}")
        evr.bold = True
        evr.font.size = Pt(11)
        evr.font.color.rgb = _NAVY

    doc.add_paragraph()

    # ── Numbered sections ──────────────────────────────────────────────────
    for i, section in enumerate(sections, start=1):
        hp = doc.add_paragraph()
        hp.paragraph_format.space_before = Pt(14)
        hp.paragraph_format.space_after = Pt(3)
        hr = hp.add_run(f"{i}.  {section.get('heading', f'Section {i}')}")
        hr.bold = True
        hr.font.size = Pt(15)
        hr.font.color.rgb = _NAVY

        body = section.get("body", "")
        for block in body.split("\n\n"):
            block = block.strip()
            if not block:
                continue
            lines = block.split("\n")
            if all(line.lstrip().startswith(("-", "•", "*")) for line in lines if line.strip()):
                for line in lines:
                    text = line.lstrip().lstrip("-•*").strip()
                    if text:
                        doc.add_paragraph(text, style="List Bullet")
            else:
                doc.add_paragraph(block.replace("\n", " "))

        # Structured tables — this is what makes a proposal executable
        for tbl in section.get("tables") or []:
            headers = [h for h in (tbl.get("headers") or []) if str(h).strip()]
            rows = tbl.get("rows") or []
            if not headers or not rows:
                continue
            if tbl.get("title"):
                cap = doc.add_paragraph()
                cap.paragraph_format.space_before = Pt(6)
                cap_run = cap.add_run(str(tbl["title"]))
                cap_run.bold = True
                cap_run.font.size = Pt(10)
                cap_run.font.color.rgb = _GOLD

            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            # header row — navy fill, white bold text
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                _shade(cell, _NAVY_HEX)
                cell.text = ""
                run = cell.paragraphs[0].add_run(str(h))
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # body rows — zebra striping
            for ridx, row in enumerate(rows):
                cells = table.add_row().cells
                if ridx % 2 == 1:
                    for c in cells:
                        _shade(c, _ZEBRA)
                for j in range(len(headers)):
                    cells[j].text = ""
                    run = cells[j].paragraphs[0].add_run(str(row[j]) if j < len(row) else "")
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = _INK
            doc.add_paragraph()

    # ── Footer note (deposit clause is business law R2) ────────────────────
    _rule(doc, "B88A2A")
    note = doc.add_paragraph()
    note_run = note.add_run(
        "Commercial terms: a 50% deposit is required before work commences (ZYNTH financial "
        "law). Figures use the Myanmar market exchange rate (sell-side), reviewed at contract "
        "date. Vendor rates tagged 'Indicative' are confirmed at contracting; 'Verified' rates "
        "are current in ZYNTH's supplier database."
    )
    note_run.font.size = Pt(8.5)
    note_run.font.color.rgb = _GREY

    _OUT_DIR.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M")
    path = _OUT_DIR / f"{_safe_filename(title)}_{stamp}.docx"
    doc.save(str(path))
    return path
