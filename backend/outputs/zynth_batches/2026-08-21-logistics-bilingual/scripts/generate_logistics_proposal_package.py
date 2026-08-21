from pathlib import Path
import json, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-logistics-bilingual')
CAMPAIGNS=json.loads((ROOT/'data/campaigns.json').read_text(encoding='utf-8'))['campaigns']
COMMERCIALS=json.loads((ROOT/'data/commercials.json').read_text(encoding='utf-8'))['commercials']
PROPOSALS=ROOT/'proposals'; STORYBOARDS=ROOT/'commercial_storyboards'; PROPOSALS.mkdir(exist_ok=True); STORYBOARDS.mkdir(exist_ok=True)
NAVY='173A5E'; TEAL='1A6A68'; GOLD='BE8834'; PALE='EAF1F6'; INK='1F2933'; GREY='6B7280'
SOURCES={
'LOG-S01':('Myanmar Logistics Institute / MIFFA','https://myanmarlogisticsinstitute.com/'),
'LOG-S02':('EuroCham Myanmar — Logistics Briefing','https://eurocham-myanmar.org/events/logistics-briefing-current-global-trade-transport-insights-for-businesses-in-myanmar/'),
'LOG-S03':('Myanma Port Authority — Annual Trade Overview','https://www.mpa.gov.mm/annual-trade-overview/'),
'LOG-S04':('Myanma Port Authority — Port Master Plan Project','https://www.mpa.gov.mm/development_projects/project-for-formulation-of-port-master-plan/'),
'LOG-S05':('Global New Light of Myanmar — association reconstitution discovery result','https://www.gnlm.com.mm/miffa-reconstituted-as-myanmar-logistics-association-mla/'),
'LOG-S06':('Myanmar Government — Upcoming Holidays','https://myanmar.gov.mm/upcoming-holidays/')}

def shade(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def cell(cell,text,bold=False,color=INK,size=8.2):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(text)); r.bold=bold; r.font.name='Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar'); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def style(doc):
    s=doc.sections[0]; s.top_margin=Inches(.55); s.bottom_margin=Inches(.55); s.left_margin=Inches(.6); s.right_margin=Inches(.6)
    for n in ['Normal','Title','Heading 1','Heading 2']:
        x=doc.styles[n]; x.font.name='Noto Sans Myanmar'; x._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar')
    doc.styles['Normal'].font.size=Pt(9); doc.styles['Normal'].font.color.rgb=RGBColor.from_string(INK)
    doc.styles['Title'].font.size=Pt(23); doc.styles['Title'].font.color.rgb=RGBColor.from_string(NAVY)
    doc.styles['Heading 1'].font.size=Pt(14); doc.styles['Heading 1'].font.color.rgb=RGBColor.from_string(TEAL)
    doc.styles['Heading 2'].font.size=Pt(10.5); doc.styles['Heading 2'].font.color.rgb=RGBColor.from_string(NAVY)
def title(doc,en,my,label,ident):
    p=doc.add_paragraph(); r=p.add_run('ZYNTH'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(GOLD); p.add_run('  |  Myanmar-First Bilingual Production System').font.size=Pt(8.5)
    doc.add_paragraph(en,style='Title'); p=doc.add_paragraph(); r=p.add_run(my); r.font.name='Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar'); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(TEAL)
    p=doc.add_paragraph(); r=p.add_run(f'{label}  •  {ident}  •  Logistics / Myanmar'); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GREY)
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.LEFT; shade(t.cell(0,0),PALE); cell(t.cell(0,0),'Client-ready planning proposal. All client, customer, carrier, fleet, route, port, warehouse, data, price, capacity, safety, permit, rights, partner and performance elements are proposed/TBC until written approval.',True,NAVY,9)
def heading(doc,en,my):
    doc.add_paragraph(en,style='Heading 1'); p=doc.add_paragraph(); r=p.add_run(my); r.italic=True; r.font.name='Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar'); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(GREY)
def kv(doc,pairs):
    t=doc.add_table(rows=0,cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for k,v in pairs:
        row=t.add_row().cells; shade(row[0],PALE); cell(row[0],k,True,NAVY); cell(row[1],v)
    doc.add_paragraph()
def table(doc,headers,rows,font=7.8):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers): shade(t.rows[0].cells[i],NAVY); cell(t.rows[0].cells[i],h,True,'FFFFFF',font)
    for n,rowdata in enumerate(rows):
        row=t.add_row().cells
        for i,v in enumerate(rowdata):
            if n%2==0: shade(row[i],'F8FAFC')
            cell(row[i],v,False,INK,font)
    doc.add_paragraph(); return t
def bullets(doc,items):
    for x in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.add_run(x)
def mmk(x): return f'MMK {x/1000000:.1f}m'
def slug(s): return re.sub(r'[^a-z0-9]+','-',s.lower()).strip('-')

def campaign_doc(c):
    d=Document(); style(d); title(d,c['shortName'],c['titleMm'],'CAMPAIGN / ACTIVATION PROPOSAL',c['id'])
    heading(d,'1. Commercial Brief','စီးပွားရေးဆိုင်ရာ အကျဉ်းချုပ်')
    kv(d,[('Format',c['format']),('Commercial tension',c['commercialTension']),('Objective',c['objective']),('Audience',c['audience']),('Creative territory',c['creativeTerritory']),('Seasonal / date logic',c['seasonalLogic'])])
    heading(d,'2. Conversion Experience','ပြောင်းလဲမှု အတွေ့အကြုံ')
    kv(d,[('Conversion mechanism',c['conversionMechanism']),('Experience flow','1) Notice the decision tension  2) Work through a controlled fictional scenario  3) Identify a next question and accountable owner  4) Give consent for a client-approved discovery conversation.')])
    heading(d,'3. Proposed Workplan','အဆိုပြုလုပ်ငန်းစဉ်')
    table(d,['Stage','Indicative timing','ZYNTH output','Client decision required'],[
        ('G0–G1: Brief & evidence','Working days 1–2','Evidence register, audience hypothesis, operational/claim boundary','Approve objective, data boundary and factual service/product context'),
        ('G2–G3: Strategy & creative','Working days 3–5','Journey, creative route, experience/content plan','Select route and priority conversion action'),
        ('G4: Feasibility','Working days 6–8','Safety, route/site, privacy, rights, staffing and risk register','Name accountable owners and clear go/no-go conditions'),
        ('G5: Production readiness','Working days 9–12','Run-of-show, asset list, training, capture and escalation plan','Approve final assets, rules, terms and operations'),
        ('G6–G7: Live & learn','Client/site-confirmed','Dashboard check, daily log, decision record and learning note','Approve optimisation or close-out')],7.7)
    heading(d,'4. Budget Envelope','ဘတ်ဂျက် စီမံကိန်း')
    b=c['budgetMMK']; rec=b['recommended']; splits=[('Strategy, creative & project direction',.15),('Experience/production operations',.34),('Staffing, training & facilitation',.18),('Data, CRM & reporting',.12),('Safety, route/site and guest operations',.11),('Contingency',.10)]
    table(d,['Planning scenario','Non-tax envelope','Use'],[('Lean',mmk(b['lean']),'Minimum viable controlled pilot'),('Recommended',mmk(rec),'Balanced production, operations and learning capacity'),('Flagship',mmk(b['flagship']),'Expanded route, content and operational capacity')])
    table(d,['Recommended budget workstream','Indicative share','Planning amount'],[(n,f'{s:.0%}',mmk(rec*s)) for n,s in splits])
    heading(d,'5. Measurement and Commercial Scenarios','စောင့်ကြည့်တိုင်းတာမှုနှင့် Scenario')
    table(d,['Scenario','Directional outcome'],[(k.title(),v) for k,v in c['planningOutcomes'].items()])
    kv(d,[('Primary KPIs','; '.join(c['primaryKpis'])),('Data rule','Capture only client-approved, consented data needed for a stated follow-up. Do not display or collect real shipment, customer, employee, route, fleet, cargo or commercial-confidential data.'),('Scenario rule','Planning assumptions only. Replace with client funnel, margin, service capacity, route, inventory and attribution data before commercial sign-off.')])
    heading(d,'6. Mandatory Preflight','မဖြစ်မနေ အတည်ပြုရမည့်အချက်များ')
    bullets(d,c['requiredApprovals']+['No price, transit time, delivery, capacity, customs, safety, security, tracking, emissions, route, warehouse, carrier, port, employment, service-level or partnership claim may be used without written client/legal approval.','All third-party visual, location, vehicle, person, uniform, system-interface, music, talent and photography rights must be cleared before production.'])
    heading(d,'7. Evidence and Use Limits','အထောက်အထားနှင့် အသုံးပြုခွင့် ကန့်သတ်ချက်')
    table(d,['ID','Public source','URL'],[(sid,SOURCES[sid][0],SOURCES[sid][1]) for sid in c['sourceUse']],7.1)
    d.add_paragraph('Evidence use: sources provide public logistics, trade, transport, capability and calendar context only. They do not validate any future commercial partnership, venue, route, carrier, capacity, performance, attendance, lead or sales outcome.')
    d.add_paragraph('Prepared by Manus AI for ZYNTH. Planning document only; not a contract, supplier quotation, legal advice, safety clearance, route plan or operational approval.',style='Caption')
    return d

def commercial_doc(c):
    d=Document(); style(d); title(d,c['titleEn'],c['titleMm'],'STANDALONE COMMERCIAL + STORYBOARD DIRECTION',c['id'])
    heading(d,'1. Film Brief','ဖիլմအကျဉ်းချုပ်')
    kv(d,[('Format',c['format']),('Linked campaign',c['linkedCampaign']),('Commercial tension',c['tension']),('Objective',c['objective']),('Creative territory',c['territory']),('Visual style',c['visualStyle']),('Storyboard status',c['storyboardStatus'])])
    heading(d,'2. 12-Frame Detailed Storyboard','အသေးစိတ် Storyboard ၁၂ ခန်း')
    rows=[(f"{x['frame']:02d} | {x['duration']}",x['beat'],x['visual'],x['camera'],x['sound'],x['onScreen']) for x in c['storyboard']]
    table(d,['Frame / time','Beat','Visual action','Camera / lens / movement','Sound / ambience','On-screen copy'],rows,6.2)
    heading(d,'3. Production & Claims Gate','ထုတ်လုပ်မှုနှင့် Claim အတည်ပြုချက်')
    bullets(d,['This is a creative treatment and storyboard direction, not a final shoot script, real-route representation, service claim, talent confirmation or costed production bid.','Lock the final script only after client factual service/product information, brand/legal review, location/site safety, privacy/security, talent, usage rights, music, subtitle language, system-interface clearance and end-card CTA are approved.','Prepare master, vertical, square and short cutdown needs only after confirmed distribution and media plan approval.','Use only consented real people and client-approved controlled settings; do not film live driving, live dispatch, real cargo, customer data, restricted zones or operational activity without specialist written clearance.'])
    heading(d,'4. Production Planning Envelope','ထုတ်လုပ်မှု စီမံကိန်း')
    table(d,['Workstream','Planning instruction'],[('Development','Confirm factual proposition, audience, CTA, route and claim boundary.'),('Pre-production','Approve cast, location/site, props, uniforms, mock interfaces, rights, safety and access protocol.'),('Production','Use a controlled set plan, simulated data, clear zones, client sign-off and daily asset backup.'),('Post-production','Build approved bilingual subtitles, versioning, music/licensing record, legal card and client review gates.'),('Delivery','Release only approved masters/cutdowns with an archive link and version status in the Master Tracker.')])
    heading(d,'5. Evidence Boundary','အထောက်အထား ကန့်သတ်ချက်')
    d.add_paragraph('The film uses the logistics research register only as broad context. It does not represent any route, port, warehouse, carrier, vehicle, customer, worker, authority, association, service level, performance metric, location or partner as confirmed. All such details remain proposed/TBC.')
    d.add_paragraph('Prepared by Manus AI for ZYNTH. Creative planning document only; not a shooting permit, service contract, legal clearance, safety clearance or operational plan.',style='Caption')
    return d

portfolio=['# ZYNTH Logistics Two-Hour Batch — Portfolio Overview','', '> **Batch:** `ZYNTH-20260821-LOGISTICS-BILINGUAL`  |  **Industry:** Logistics and supply chain  |  **Status:** Research-grounded concepts; all client, commercial and production elements proposed/TBC.','','## Ten Campaign / Activation Proposals','','| ID | Campaign | Format | Conversion mechanism | Recommended planning envelope |','|---|---|---|---|---:|']
for c in CAMPAIGNS: portfolio.append(f"| {c['id']} | {c['shortName']} | {c['format']} | {c['conversionMechanism']} | {mmk(c['budgetMMK']['recommended'])} |")
portfolio += ['','## Ten Separate Commercial / Storyboard Concepts','','| ID | Commercial | Format | Linked campaign | Storyboard frames |','|---|---|---|---|---:|']
for c in COMMERCIALS: portfolio.append(f"| {c['id']} | {c['titleEn']} | {c['format']} | {c['linkedCampaign']} | {len(c['storyboard'])} |")
portfolio += ['','## Research Sources','']+[f'- **{k}:** [{v[0]}]({v[1]})' for k,v in SOURCES.items()]+['','## Operating Note','','Campaign and commercial tracks are intentionally independent. Each campaign has a `CMP` ID; each commercial has a `COM` ID. Link them only where strategically useful, then keep budget, rights, production, approval and learning records separate in the live ZYNTH Master Tracker.']
(ROOT/'ZYNTH_LOGISTICS_Batch_Portfolio_Overview.md').write_text('\n'.join(portfolio)+'\n',encoding='utf-8')
for i,c in enumerate(CAMPAIGNS,1):
    stem=f'{i:02d}_{slug(c["shortName"])}_Campaign_Proposal_MM-Bilingual'; campaign_doc(c).save(PROPOSALS/(stem+'.docx'))
    md=['# '+c['shortName'],'',c['titleMm'],'',f'**ID:** {c["id"]}  ',f'**Format:** {c["format"]}  ','','## Commercial Brief','',f'**Tension:** {c["commercialTension"]}','',f'**Objective:** {c["objective"]}','',f'**Creative territory:** {c["creativeTerritory"]}','', '## Conversion Experience','',c['conversionMechanism'],'', '## Budget (MMK, non-tax planning envelope)','',f"Lean: {mmk(c['budgetMMK']['lean'])} | Recommended: {mmk(c['budgetMMK']['recommended'])} | Flagship: {mmk(c['budgetMMK']['flagship'])}",'','## Primary KPIs','']+['- '+x for x in c['primaryKpis']]+['','## Required Approvals','']+['- '+x for x in c['requiredApprovals']]+['','> All client, venue/site, route, customer, data, safety, claims, permits, rights and commercial details remain proposed/TBC. Scenario outcomes are planning assumptions, not guarantees.']
    (PROPOSALS/(stem+'.md')).write_text('\n'.join(md)+'\n',encoding='utf-8')
for i,c in enumerate(COMMERCIALS,1):
    stem=f'{i:02d}_{slug(c["titleEn"])}_Commercial_Storyboard_MM-Bilingual'; commercial_doc(c).save(STORYBOARDS/(stem+'.docx'))
    md=['# '+c['titleEn'],'',c['titleMm'],'',f'**ID:** {c["id"]}  ',f'**Format:** {c["format"]}  ',f'**Linked campaign:** {c["linkedCampaign"]}  ','','## Film Brief','',f'**Tension:** {c["tension"]}','',f'**Objective:** {c["objective"]}','',f'**Creative territory:** {c["territory"]}','', '## 12-Frame Storyboard','', '| # | Time | Beat | Visual action | Camera | Sound | Text |','|---:|---|---|---|---|---|---|']
    for x in c['storyboard']: md.append('| '+str(x['frame'])+' | '+' | '.join(str(x[k]).replace('|','/') for k in ['duration','beat','visual','camera','sound','onScreen'])+' |')
    md += ['','## Production Gate','', '> This is a creative storyboard direction only. Script, cast, locations, service/product facts, claims, system interfaces, rights, music, subtitles, safety, end-card CTA and release plan require written client approval.']
    (STORYBOARDS/(stem+'.md')).write_text('\n'.join(md)+'\n',encoding='utf-8')
print(f'Created {len(CAMPAIGNS)} campaign Word + Markdown proposals and {len(COMMERCIALS)} commercial Word + Markdown storyboards.')
