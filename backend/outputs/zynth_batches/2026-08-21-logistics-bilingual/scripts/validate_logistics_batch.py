from pathlib import Path
import json
from docx import Document

ROOT=Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-logistics-bilingual')
CAM=json.loads((ROOT/'data/campaigns.json').read_text(encoding='utf-8'))['campaigns']
COM=json.loads((ROOT/'data/commercials.json').read_text(encoding='utf-8'))['commercials']
passed=[]; issues=[]
def check(ok,name,detail=''):(passed if ok else issues).append((name,detail))
def doc_text(p):
    d=Document(p); out=[x.text for x in d.paragraphs]
    for t in d.tables:
        for row in t.rows: out.extend(c.text for c in row.cells)
    return '\n'.join(out)
check(len(CAM)==10,'Campaign count',str(len(CAM))); check(len(COM)==10,'Commercial count',str(len(COM))); check(sum(len(x['storyboard']) for x in COM)==120,'Storyboard-frame count',str(sum(len(x['storyboard']) for x in COM)))
for field in ['format','commercialTension','conversionMechanism','creativeTerritory','seasonalLogic']:
    check(len(set(x[field] for x in CAM))==10,f'Campaign unique {field}',f'{len(set(x[field] for x in CAM))}/10')
check(len(set(x['format'] for x in COM))==10,'Commercial unique format',f'{len(set(x["format"] for x in COM))}/10'); check(len(set(x['territory'] for x in COM))==10,'Commercial unique territory',f'{len(set(x["territory"] for x in COM))}/10')
cp=sorted((ROOT/'proposals').glob('*.docx')); cm=sorted((ROOT/'proposals').glob('*.md')); sp=sorted((ROOT/'commercial_storyboards').glob('*.docx')); sm=sorted((ROOT/'commercial_storyboards').glob('*.md'))
for name,found in [('Campaign Word documents',cp),('Campaign Markdown documents',cm),('Commercial Word documents',sp),('Commercial Markdown documents',sm)]:check(len(found)==10,name,str(len(found)))
for p in cp:
    text=doc_text(p)
    for token in ['Commercial Brief','Conversion Experience','Budget Envelope','Mandatory Preflight','Evidence and Use Limits']: check(token in text,f'Campaign heading {token}',p.name)
    check('proposed/TBC' in text,f'Campaign boundary disclaimer',p.name)
for p in sp:
    text=doc_text(p)
    for token in ['Film Brief','12-Frame Detailed Storyboard','Production & Claims Gate','Evidence Boundary']:check(token in text,f'Commercial heading {token}',p.name)
    check('creative treatment' in text.lower(),f'Commercial treatment disclaimer',p.name)
for p in cm+sm:
    text=p.read_text(encoding='utf-8').lower()
    for bad in ['official partner','confirmed sponsor','award-winning','guaranteed delivery','guaranteed capacity','transit time guarantee','safe for']:
        check(bad not in text,f'No unverified phrase: {bad}',p.name)
lines=['# ZYNTH Logistics Batch Validation','', '**Batch:** `ZYNTH-20260821-LOGISTICS-BILINGUAL`  ',f'**Pass checks:** {len(passed)}  ',f'**Exceptions:** {len(issues)}  ','', '## Result','', '**PASS**' if not issues else '**REVIEW REQUIRED**','', '## Checks Passed','']+[f'- PASS — {a}: {b}' for a,b in passed]
if issues: lines += ['','## Exceptions','']+[f'- REVIEW — {a}: {b}' for a,b in issues]
(ROOT/'validation'/'batch_validation.md').write_text('\n'.join(lines)+'\n',encoding='utf-8')
print(f'PASS={len(passed)} EXCEPTIONS={len(issues)}')
if issues: raise SystemExit(1)
