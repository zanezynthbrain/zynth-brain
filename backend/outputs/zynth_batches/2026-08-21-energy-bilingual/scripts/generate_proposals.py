from __future__ import annotations

import json
import math
from pathlib import Path
from datetime import datetime, timezone

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-energy-bilingual')
DATA = ROOT / 'data' / 'concepts.json'
DOCS = ROOT / 'proposals'
BATCH = 'ZYNTH-20260821-ENERGY-BILINGUAL'
VERSION = 'v1.0 — planning draft'
GENERATED = '21 August 2026 (UTC)'
SOURCE_LOG = [
    ('S01', 'Reuters (14 Nov 2025)', 'Myanmar solar adoption amid unreliable supply; context only.', 'https://www.reuters.com/sustainability/climate-energy/war-torn-myanmar-embraces-solar-tackle-power-crisis-2025-11-14/'),
    ('S02', 'Xinhua (18 Jan 2026)', 'Dated 2026 Yangon solar-expo precedent; no future rights or repeat event presumed.', 'https://english.news.cn/20260118/765dc93ac49a481d8ffbcc1972c01516/c.html'),
    ('S03', 'Smart Power Myanmar (accessed 21 Aug 2026)', 'Published financing-to-installation journey; supplier self-reported case studies not used as sector benchmark.', 'https://www.smartpowermyanmar.org/'),
    ('S04', 'Global New Light of Myanmar (25 Jul 2026)', 'Published state context on solar use and technology-management discussion; not an approval.', 'https://www.gnlm.com.mm/efficiently-generate-solar-energy-for-the-state/'),
    ('S05', 'Myanmar Ministry of Foreign Affairs / Myanmar National Portal', 'Official 2026 Q4 holiday dates used only as timing guardrails.', 'https://www.mofa.gov.mm/about-myanmar/public-holidays/'),
    ('S06', 'Livoltek YouTube expo video (reviewed 21 Aug 2026)', 'Visual-format reference only; conflicting date signals excluded from factual planning.', 'https://www.youtube.com/watch?v=9-S7urcEX2U'),
]

COLORS = {
    'navy': '0E3150', 'gold': 'E2A744', 'ice': 'EFF5F8', 'ink': '17212B',
    'muted': '5D6A75', 'white': 'FFFFFF', 'green': '2F7B61', 'red': 'A53A38'
}

MY = {
    'proposal': 'အဆိုပြုချက်',
    'planning': 'Planning Document',
    'overview': 'အနှစ်ချုပ်',
    'strategic': 'မဟာဗျူဟာ အခြေခံ',
    'insight': 'Audience Insight',
    'behaviour': 'လိုချင်သော အပြုအမူပြောင်းလဲမှု',
    'journey': 'Audience Journey',
    'plan': 'Experience / Campaign Plan',
    'season': 'အချိန်အခါနှင့် Calendar Logic',
    'content': 'Content System',
    'talent': 'Talent Logic',
    'production': 'Production Requirements',
    'budget': 'MMK Budget Packages',
    'roi': 'Commercial Scenario နှင့် Break-even Logic',
    'kpi': 'KPI Scorecard',
    'risk': 'Risk / Compliance',
    'workflow': 'Workflow, Gates နှင့် Approvals',
    'treatment': 'Commercial Video Treatment နှင့် Storyboard',
    'design': 'Design / Experience Package',
    'sources': 'Source Log',
    'ask': 'Approval Ask',
    'disclaimer': 'အရေးကြီးသော Planning သတိပြုရန်',
}


def mmk(value: int) -> str:
    return f'MMK {value:,.0f}'


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold=False, color=None, size=8.5, font='Aptos'):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = font
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('ZYNTH • CONFIDENTIAL PLANNING DOCUMENT • ')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(COLORS['muted'])
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def configure_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Inches(0.62)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.68)
    sec.right_margin = Inches(0.68)
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Aptos'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = RGBColor.from_string(COLORS['ink'])
    for style_name, size, color in [('Title', 28, COLORS['navy']), ('Heading 1', 17, COLORS['navy']), ('Heading 2', 12, COLORS['navy'])]:
        st = styles[style_name]
        st.font.name = 'Aptos Display'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
    footer = sec.footer
    add_page_number(footer.paragraphs[0])


def add_label(doc, text, my=False):
    p = doc.add_paragraph()
    p.space_after = Pt(3)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(COLORS['gold'])
    r.font.name = 'Noto Sans Myanmar' if my else 'Aptos'
    return p


def add_title(doc, title, subtitle='', my=False):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(23)
    r.font.name = 'Noto Sans Myanmar' if my else 'Aptos Display'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    r.font.color.rgb = RGBColor.from_string(COLORS['navy'])
    if subtitle:
        q = doc.add_paragraph()
        q.paragraph_format.space_after = Pt(9)
        rr = q.add_run(subtitle)
        rr.font.name = 'Noto Sans Myanmar' if my else 'Aptos'
        rr._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
        rr.font.size = Pt(11)
        rr.font.color.rgb = RGBColor.from_string(COLORS['muted'])


def heading(doc, title, my=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(COLORS['navy'])
    r.font.name = 'Noto Sans Myanmar' if my else 'Aptos Display'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    return p


def para(doc, text, my=False, italic=False, color=None, size=9.2):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Noto Sans Myanmar' if my else 'Aptos'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def bullets(doc, items, my=False):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(item)
        r.font.size = Pt(9)
        r.font.name = 'Noto Sans Myanmar' if my else 'Aptos'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')


def table(doc, headers, rows, widths=None, my=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        shade(hdr.cells[i], COLORS['navy'])
        set_cell_text(hdr.cells[i], header, bold=True, color=COLORS['white'], size=8.2, font='Noto Sans Myanmar' if my else 'Aptos')
        set_cell_margins(hdr.cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, size=8.1, font='Noto Sans Myanmar' if my else 'Aptos')
            set_cell_margins(cells[i])
            if len(t.rows) % 2 == 1:
                shade(cells[i], 'F7FAFC')
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return t


def add_cover(doc, c, my=False):
    add_label(doc, 'ZYNTH TWO-HOUR BILINGUAL PRODUCTION PLAYBOOK' if not my else 'ZYNTH Myanmar-First Agency Production Batch', my)
    title = c['title_my'] if my else c['title_en']
    subtitle = c['title_en'] if my else c['title_my']
    add_title(doc, title, subtitle, my)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(15)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(c['tagline_my'] if my else c['tagline'])
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = 'Noto Sans Myanmar' if my else 'Aptos Display'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    r.font.color.rgb = RGBColor.from_string(COLORS['gold'])
    cover_rows = [
        ('Industry / Industry Code', 'Energy and solar adoption / energy'),
        ('Proposal format' if not my else 'အဆိုပြုထားသည့် Format', c['format'] if not my else c['format']),
        ('Batch' if not my else 'Batch Code', BATCH),
        ('Document version' if not my else 'Version', VERSION),
        ('Prepared' if not my else 'ပြင်ဆင်သည့်ရက်', GENERATED),
        ('Source status' if not my else 'Source Status', 'New ZYNTH planning concept; no client, venue, partner or supplier appointment is implied.' if not my else 'ZYNTH planning concept အသစ်ဖြစ်ပြီး client၊ venue၊ partner သို့မဟုတ် supplier appointment မည်သည့်အရာကိုမျှ မဆိုလိုပါ။')
    ]
    table(doc, ['DOCUMENT CONTROL' if not my else 'DOCUMENT CONTROL', 'DETAIL' if not my else 'အချက်အလက်'], cover_rows, [2.0, 4.7], my)
    doc.add_paragraph()
    disclaimer = ('This is a planning document, not a supplier quotation, financial guarantee, legal approval, permit confirmation, engineering drawing, performance warranty or partnership announcement. All named places, talent, production houses, rights, permits, engineering specifications and media inventory remain proposed/TBC until independently verified and client-approved.')
    if my:
        disclaimer = ('ဤစာတမ်းသည် Planning Document သာဖြစ်ပြီး supplier quotation၊ financial guarantee၊ legal approval၊ permit confirmation၊ engineering drawing၊ performance warranty သို့မဟုတ် partnership announcement မဟုတ်ပါ။ ဖော်ပြထားသော place၊ talent၊ production house၊ rights၊ permit၊ engineering specification နှင့် media inventory အားလုံးသည် သီးခြားစစ်ဆေးပြီး client approval ရရှိမှသာ အတည်ဖြစ်မည်ဖြစ်ကာ လက်ရှိတွင် proposed/TBC ဖြစ်ပါသည်။')
    para(doc, disclaimer, my, italic=True, color=COLORS['muted'], size=8.6)
    doc.add_page_break()


def add_strategy(doc, c, my=False):
    heading(doc, MY['overview'] if my else 'Executive Overview', my)
    overview = (f"{c['format']} built for a client-selected energy or solar brand. The concept turns the commercial tension—{c['commercial_tension'].lower()}—into a consented next step: {c['conversion']}")
    if my:
        overview = (f"Client မှရွေးချယ်မည့် energy သို့မဟုတ် solar brand အတွက် {c['format']} ကို အဆိုပြုထားပါသည်။ {c['tension_my']} ဟူသော commercial tension ကို {c['conversion_my']} အဖြစ် ပြောင်းလဲပေးရန် ရည်ရွယ်ပါသည်။")
    para(doc, overview, my)
    table(doc,
          ['STRATEGIC FIELD' if not my else 'မဟာဗျူဟာအချက်', 'PROPOSAL RESPONSE' if not my else 'အဆိုပြုချက်'],
          [
              ('Commercial tension' if not my else 'Commercial Tension', c['commercial_tension'] if not my else c['tension_my']),
              ('Creative territory' if not my else 'Creative Territory', c['territory']),
              ('Behaviour change' if not my else 'လိုချင်သောပြောင်းလဲမှု', c['behaviour_change'] if not my else c['behaviour_my']),
              ('Conversion mechanism' if not my else 'Conversion Mechanism', c['conversion'] if not my else c['conversion_my']),
              ('Core message' if not my else 'အဓိက Message', c['message'] if not my else c['message_my']),
              ('CTA example' if not my else 'CTA ဥပမာ', c['cta'] if not my else c['cta_my']),
          ], [1.65, 5.05], my)
    heading(doc, MY['insight'] if my else 'Audience Insight', my)
    para(doc, c['insight_my'] if my else c['insight'], my)
    heading(doc, MY['behaviour'] if my else 'Audience Behaviour Change', my)
    para(doc, c['behaviour_my'] if my else c['behaviour_change'], my)
    heading(doc, MY['journey'] if my else 'Audience Journey', my)
    journey = [
        ('1. See' if not my else '၁။ မြင်တွေ့', 'Recognises a familiar interruption or decision friction.' if not my else 'ရင်းနှီးနေသော interruption သို့မဟုတ် decision friction ကို သတိပြုမိသည်။'),
        ('2. Understand' if not my else '၂။ နားလည်', c['message'] if not my else c['message_my']),
        ('3. Experience' if not my else '၃။ အတွေ့အကြုံ', 'Receives a guided, safe proof moment with claims bounded by approvals.' if not my else 'Approval ဖြင့် ကန့်သတ်ထားသော claim များအတွင်း guided နှင့် safe proof moment ကို ရရှိသည်။'),
        ('4. Choose' if not my else '၄။ ရွေးချယ်', c['cta'] if not my else c['cta_my']),
        ('5. Continue' if not my else '၅။ ဆက်လက်', 'A named owner follows up only where consent and client governance permit.' if not my else 'Consent နှင့် client governance ခွင့်ပြုသည့်နေရာတွင်သာ အမည်တပ် owner က follow-up ပြုလုပ်သည်။'),
    ]
    table(doc, ['STAGE' if not my else 'အဆင့်', 'EXPERIENCE' if not my else 'Experience'], journey, [1.35, 5.35], my)
    heading(doc, MY['season'] if my else 'Seasonal / Special-Day Rationale', my)
    para(doc, c['seasonal_my'] if my else c['seasonal'], my)
    note = 'The only dated Q4 public calendar references used in this proposal are official timing cues; they create no event right, sponsor entitlement, public-activation approval or permit.'
    if my:
        note = 'ဤအဆိုပြုချက်တွင် အသုံးပြုထားသည့် Q4 public calendar ရက်စွဲများသည် official timing cue များသာဖြစ်ပြီး event right၊ sponsor entitlement၊ public activation approval သို့မဟုတ် permit မပေးပါ။'
    para(doc, note, my, italic=True, color=COLORS['muted'], size=8.5)
    doc.add_page_break()


def add_plan(doc, c, my=False):
    heading(doc, MY['plan'] if my else 'Experience / Campaign Plan', my)
    if my:
        planned = [
            'Client-approved audience list သို့မဟုတ် consented media audience ကိုသာ အသုံးပြုမည်။',
            'Message → proof moment → consented next step အစဉ်ဖြင့် audience journey ကို တည်ဆောက်မည်။',
            'Sale, saving, uptime, technical output နှင့် availability claim များကို client legal/compliance approval မတိုင်မီ မထုတ်ပြန်ပါ။'
        ]
    else:
        planned = c['plan']
    bullets(doc, planned, my)
    heading(doc, MY['content'] if my else 'Content System', my)
    para(doc, c['content_system'] if not my else 'Hero content တစ်ခု၊ short-form cutdown များ၊ bilingual CTA asset များနှင့် consented follow-up sequence ကို တစ်စနစ်တည်းဖြင့် ထုတ်လုပ်ပါမည်။ ' + c['content_system'], my)
    content_rows = [
        ('Hero asset' if not my else 'Hero Asset', '1 × 45–60 second commercial video with 16:9, 9:16 and 1:1 masters.' if not my else '16:9၊ 9:16 နှင့် 1:1 master ပါဝင်သော 45–60 second commercial video တစ်ခု။'),
        ('Explainers' if not my else 'Explainer များ', '3–6 × 10–20 second proof, question or journey cutdowns.' if not my else 'Proof၊ question သို့မဟုတ် journey အတွက် 10–20 second cutdown 3–6 ခု။'),
        ('Conversion assets' if not my else 'Conversion Asset များ', 'Landing/registration/CRM or event follow-up assets subject to data and legal review.' if not my else 'Data နှင့် legal review အောက်ရှိ landing/registration/CRM သို့မဟုတ် event follow-up asset များ။'),
        ('Measurement assets' if not my else 'Measurement Asset များ', 'UTM/QR/event tags, consent log and client-approved dashboard fields.' if not my else 'UTM/QR/event tag၊ consent log နှင့် client-approved dashboard field များ။'),
    ]
    table(doc, ['LAYER' if not my else 'အလွှာ', 'DELIVERABLE' if not my else 'Deliverable'], content_rows, [1.75, 5.0], my)
    heading(doc, MY['talent'] if my else 'Talent Logic', my)
    para(doc, c['talent_logic'] if not my else 'Talent အတွက် ' + c['talent_logic'] + ' Casting approval၊ conduct protocol၊ language suitability၊ usage territory/term နှင့် release များကို shoot မတိုင်မီ အတည်ပြုရပါမည်။', my)
    heading(doc, MY['production'] if my else 'Production Requirements', my)
    para(doc, c['production'] if not my else 'Production လိုအပ်ချက် — ' + c['production'] + ' Supplier engineering၊ venue assessment၊ risk assessment နှင့် client approval မပြီးမချင်း design ကို build-ready အဖြစ် မယူဆရပါ။', my)
    doc.add_page_break()


def budget_rows(total):
    return [
        ('Pass-through production / media / suppliers', int(total * 0.70)),
        ('ZYNTH strategy, creative, PM and production fee', int(total * 0.20)),
        ('Contingency (uncommitted; client release required)', total - int(total * 0.70) - int(total * 0.20)),
        ('Taxes', 'Excluded / TBC at applicable rate'),
        ('Non-tax planning envelope', total),
    ]


def add_commercials(doc, c, my=False):
    heading(doc, MY['budget'] if my else 'Three Planning Budget Packages', my)
    intro = 'All MMK figures are non-binding planning envelopes, not supplier quotations. Taxes are excluded/TBC at the applicable rate; scope, timing, venue, engineering, travel, media, talent, rights, freight and permits can change the final cost.'
    if my:
        intro = 'MMK အားလုံးသည် non-binding planning envelope များသာဖြစ်ပြီး supplier quotation မဟုတ်ပါ။ Tax ကို applicable rate အတိုင်း Excluded/TBC ထားပြီး scope၊ timing၊ venue၊ engineering၊ travel၊ media၊ talent၊ rights၊ freight နှင့် permit များကြောင့် final cost ပြောင်းလဲနိုင်ပါသည်။'
    para(doc, intro, my, italic=True, color=COLORS['muted'], size=8.5)
    rows = []
    package_names = ['Lean', 'Recommended', 'Flagship']
    for name, amount in zip(package_names, c['budgets']):
        rows.append((name, mmk(amount), 'Compact proof / one core format' if name == 'Lean' else ('Balanced production, content and conversion operations' if name == 'Recommended' else 'Expanded footprint, asset suite and contingency depth')))
    if my:
        rows = [(a, b, 'Compact proof / core format တစ်ခု' if a == 'Lean' else ('Balanced production, content နှင့် conversion operations' if a == 'Recommended' else 'Expanded footprint, asset suite နှင့် contingency depth')) for a,b,_ in rows]
    table(doc, ['PACKAGE' if not my else 'Package', 'NON-TAX MMK' if not my else 'Tax မပါဝင်သော MMK', 'PLANNING INTENT' if not my else 'Planning Intent'], rows, [1.05, 1.6, 4.1], my)
    for name, amount in zip(package_names, c['budgets']):
        subrows = budget_rows(amount)
        table(doc, [name + (' BREAKDOWN' if not my else ' Breakdown'), 'MMK / STATUS' if not my else 'MMK / Status'], [(a, mmk(b) if isinstance(b, int) else b) for a,b in subrows], [4.65, 2.1], my)
    heading(doc, MY['roi'] if my else 'Scenario-Based Commercial Outcomes & Break-Even', my)
    r = c['roi']
    recommended = c['budgets'][1]
    scenario_rows = []
    for scenario, leads, rate in [('Conservative', r['conservative_leads'], r['conservative_close']), ('Base', r['base_leads'], r['base_close']), ('Upside', r['upside_leads'], r['upside_close'])]:
        expected_closes = leads * rate
        contribution = expected_closes * r['contribution_per_close']
        roi = (contribution - recommended) / recommended
        scenario_rows.append((scenario, f'{leads}', f'{rate:.0%}', f'{expected_closes:.1f}', mmk(round(contribution)), f'{roi:.0%}'))
    if my:
        scenario_rows = [('Conservative' if a=='Conservative' else ('Base' if a=='Base' else 'Upside'), b,c,d,e,f) for a,b,c,d,e,f in scenario_rows]
    table(doc, ['SCENARIO' if not my else 'Scenario', 'QUALIFIED LEADS' if not my else 'Qualified Leads', 'ASSUMED CLOSE' if not my else 'ယူဆထားသည့် Close', 'EXPECTED CLOSES' if not my else 'Expected Closes', 'INCREMENTAL CONTRIBUTION' if not my else 'Incremental Contribution', 'ROI vs Recommended' if not my else 'Recommended နှင့်နှိုင်း ROI'], scenario_rows, [0.9, 0.95, 0.8, 0.9, 1.65, 0.85], my)
    break_even = math.ceil(recommended / r['contribution_per_close'])
    msg = f"Recommended-package break-even occurs at {break_even} incremental closed deal(s) at the assumed incremental gross contribution of {mmk(r['contribution_per_close'])} per closed deal. This is a planning equation, not a revenue, performance or financial guarantee."
    if my:
        msg = f"Recommended package အတွက် break-even သည် closed deal တစ်ခုစီလျှင် ယူဆထားသော incremental gross contribution {mmk(r['contribution_per_close'])} ဖြင့် incremental closed deal {break_even} ခုရရှိသည့်အခါ ဖြစ်ပါသည်။ ၎င်းသည် planning equation သာဖြစ်ပြီး revenue၊ performance သို့မဟုတ် financial guarantee မဟုတ်ပါ။"
    para(doc, msg, my, italic=True, color=COLORS['muted'])
    doc.add_page_break()


def add_measurement(doc, c, my=False):
    heading(doc, MY['kpi'] if my else 'KPI Scorecard & Monitoring', my)
    base = c['roi']['base_leads']
    rows = [
        ('Reach / invitations' if not my else 'Reach / Invitation', 'Client-approved reach and invitation delivery', 'Platform/event evidence', 'Weekly / event close' if not my else 'အပတ်စဉ် / event close'),
        ('Engaged proof actions' if not my else 'Proof Engagement', 'Demo completion, video completion or UI progress', 'QR / analytics event', 'Daily' if not my else 'နေ့စဉ်'),
        ('Qualified leads' if not my else 'Qualified Lead', f'Base planning scenario: {base}', 'Consent log + qualification fields', 'Daily / weekly' if not my else 'နေ့စဉ် / အပတ်စဉ်'),
        ('Booked next step' if not my else 'Booked Next Step', 'Assessment, survey, workshop or decision sprint', 'CRM owner status', 'Weekly' if not my else 'အပတ်စဉ်'),
        ('Commercial outcome' if not my else 'Commercial Outcome', 'Client-defined qualified opportunity and closed-won contribution', 'Client CRM; restricted access', 'Monthly' if not my else 'လစဉ်'),
        ('Trust / safety' if not my else 'Trust / Safety', 'Claims, consent, incident and escalation log', 'PM / legal log', 'Real time' if not my else 'အချိန်နှင့်တပြေးညီ'),
    ]
    if my:
        rows = [(a,b,c,d) for a,b,c,d in rows]
    table(doc, ['KPI' if not my else 'KPI', 'PLANNING TARGET / DEFINITION' if not my else 'Planning Target / Definition', 'SOURCE' if not my else 'Source', 'CADENCE' if not my else 'Cadence'], rows, [1.25, 2.75, 1.55, 1.15], my)
    para(doc, 'No historical client conversion, margin or media-performance data was supplied. Targets are therefore working assumptions to be replaced during discovery; all personal or CRM data requires client-approved lawful basis, notice, retention and access controls.' if not my else 'Historical client conversion၊ margin သို့မဟုတ် media-performance data မပေးထားသဖြင့် target များသည် discovery အတွင်း အစားထိုးရန် working assumption များသာဖြစ်ပါသည်။ Personal သို့မဟုတ် CRM data အားလုံးအတွက် client-approved lawful basis၊ notice၊ retention နှင့် access control လိုအပ်ပါသည်။', my, italic=True, color=COLORS['muted'], size=8.5)
    heading(doc, MY['risk'] if my else 'Risk, Compliance & Safety Register', my)
    risks = [
        ('Claims / performance', 'No saving, uptime, output, financing or payback claim without substantiation and client legal approval.', 'Client legal + technical owner'),
        ('Permits / venue', 'All venue, authority, neighbourhood, site and public-activation permissions are proposed/TBC until written confirmation.', 'Client + venue / local authority'),
        ('Electrical / build safety', 'Supplier engineering, risk assessment, load plan, fire route and trained crew required before build.', 'Production house + HSE'),
        ('Privacy / CRM', 'Use consent, opt-out, limited fields, retention rule and named data owner; no contact scraping.', 'Client data owner'),
        ('Talent / rights', 'Written release, territory, term, media, music, location and voice usage rights before publication.', 'Client marketing + production'),
        ('Cultural / calendar', 'Avoid holiday assumptions and obtain cultural-sensitivity review; do not imply government affiliation.', 'Client + local reviewer'),
        ('Community safeguarding', 'Use informed consent, non-coercive participation, safeguarding protocol and transparent selection language.', 'Client CSR / safeguarding lead'),
    ]
    if my:
        risks = [
            ('Claim / Performance', 'Substantiation နှင့် client legal approval မရှိဘဲ saving, uptime, output, financing သို့မဟုတ် payback claim မပြုလုပ်ပါ။', 'Client legal + technical owner'),
            ('Permit / Venue', 'Venue, authority, neighbourhood, site နှင့် public activation permission အားလုံးသည် written confirmation မတိုင်မီ proposed/TBC ဖြစ်ပါသည်။', 'Client + venue / local authority'),
            ('Electrical / Build Safety', 'Build မတိုင်မီ supplier engineering, risk assessment, load plan, fire route နှင့် trained crew လိုအပ်ပါသည်။', 'Production house + HSE'),
            ('Privacy / CRM', 'Consent, opt-out, limited field, retention rule နှင့် named data owner ကို အသုံးပြုပါ။', 'Client data owner'),
            ('Talent / Rights', 'Publication မတိုင်မီ release, territory, term, media, music, location နှင့် voice usage rights ကို စာဖြင့်အတည်ပြုပါ။', 'Client marketing + production'),
            ('Culture / Calendar', 'Holiday assumption ကိုရှောင်ရှားပြီး cultural-sensitivity review ရယူပါ။ Government affiliation မဆိုလိုပါ။', 'Client + local reviewer'),
            ('Community Safeguarding', 'Informed consent, non-coercive participation, safeguarding protocol နှင့် transparent selection language ကို အသုံးပြုပါ။', 'Client CSR / safeguarding lead'),
        ]
    table(doc, ['RISK' if not my else 'Risk', 'CONTROL' if not my else 'Control', 'OWNER' if not my else 'Owner'], risks, [1.2, 4.25, 1.25], my)
    doc.add_page_break()


def add_design(doc, c, my=False):
    heading(doc, MY['design'] if my else 'Design / Experience Package', my)
    if c['mode'] == 'physical':
        text = ('Separate design files are supplied as: (1) labelled sketch package, source SVG + client-viewable PNG; and (2) conceptual isometric 3D-style package, source SVG + client-viewable PNG. Each includes a hero perspective, front/stage elevation, plan/top view and detail view. These are visual concepts only, not construction drawings, electrical schematics or supplier-engineered specifications.')
        if my:
            text = ('သီးခြား Design file များကို (၁) labelled sketch package—source SVG နှင့် client-viewable PNG၊ (၂) conceptual isometric 3D-style package—source SVG နှင့် client-viewable PNG အဖြစ် ထုတ်ပေးပါမည်။ Package တစ်ခုစီတွင် hero perspective၊ front/stage elevation၊ plan/top view နှင့် detail view ပါဝင်ပါသည်။ ၎င်းတို့သည် visual concept များသာဖြစ်ပြီး construction drawing၊ electrical schematic သို့မဟုတ် supplier-engineered specification မဟုတ်ပါ။')
        para(doc, text, my)
        rows = [
            ('Design format' if not my else 'Design Format', 'A3 landscape SVG source + 300 dpi PNG client preview; dimensions are approximate.', 'Concept source only; proposed/TBC'),
            ('Approx. footprint' if not my else 'ခန့်မှန်း Footprint', c['production'].split(',')[0], 'Subject to venue survey / crowd / fire route'),
            ('Materials / surfaces' if not my else 'Material / Surface', 'Powder-coated modular frame, FSC plywood or equivalent counters, tension fabric, non-slip vinyl, recyclable signage.', 'Supplier samples and structural method TBC'),
            ('Lighting / furniture' if not my else 'Lighting / Furniture', 'Low-glare LED practicals, movable stools/counters, accessible seating, labelled demonstration plinths.', 'Electrical load and fixture placement subject to engineering'),
            ('Wayfinding / zones' if not my else 'Wayfinding / Zone', 'Welcome, proof/demo, consultation, consent/registration, product/talent and exit-feedback zones.', 'Final audience flow subject to site and client approval'),
            ('Sightlines / limitations' if not my else 'Sightline / Limitation', 'Open front, 1.2m primary sightline and protected technical rear; no live operation or load representation without qualified supplier design.', 'Lead time: 4–8 weeks after approvals; complexity: medium to high'),
            ('Approvals' if not my else 'Approval များ', 'Client brand/legal, venue, HSE, electrical supplier/engineer, landlord/authority where applicable, insurance and production house.', 'Required before procurement or build'),
        ]
        if my:
            rows = [(a,b,c) for a,b,c in rows]
        table(doc, ['SPECIFICATION' if not my else 'Specification', 'DESIGN INTENT' if not my else 'Design Intent', 'STATUS / DEPENDENCY' if not my else 'Status / Dependency'], rows, [1.4, 3.6, 1.75], my)
    else:
        text = ('Separate digital design file supplied as a campaign UI/experience storyboard: mobile-first user flow, screen states, consent moments, error and accessibility notes, data-event map and CTA handoff. It is an experience concept—not a working application, data-protection assessment or production-ready technical specification.')
        if my:
            text = ('သီးခြား Digital Design file ကို campaign UI/experience storyboard အဖြစ် ထုတ်ပေးပါမည်—mobile-first user flow၊ screen state များ၊ consent moment၊ error/accessibility note၊ data-event map နှင့် CTA handoff ပါဝင်ပါသည်။ ၎င်းသည် experience concept သာဖြစ်ပြီး working application၊ data-protection assessment သို့မဟုတ် production-ready technical specification မဟုတ်ပါ။')
        para(doc, text, my)
        rows = [
            ('Design format' if not my else 'Design Format', '16:9 storyboard board and 9:16 mobile UI flow; SVG source + PNG client preview.', 'New ZYNTH concept; proposed/TBC'),
            ('Core screens' if not my else 'Core Screen', 'Entry promise, priority selector, evidence/question step, consent, booking and confirmation.', 'Final fields subject to client privacy/legal review'),
            ('Accessibility' if not my else 'Accessibility', 'Plain-language Myanmar/English copy, contrast, captions, tap targets, keyboard/readability guidance.', 'User testing and localisation QA required'),
            ('Analytics' if not my else 'Analytics', 'Consent-aware events only: view, progress, submit, booking, opt-out.', 'Tagging, CRM and retention policy TBC'),
            ('Implementation' if not my else 'Implementation', 'Responsive build, secure hosting, QA and analytics handoff require separately scoped client-approved technical work.', 'Estimated 3–6 weeks after approved requirements'),
        ]
        table(doc, ['SPECIFICATION' if not my else 'Specification', 'DESIGN INTENT' if not my else 'Design Intent', 'STATUS / DEPENDENCY' if not my else 'Status / Dependency'], rows, [1.4, 3.6, 1.75], my)
    doc.add_page_break()


def add_treatment(doc, c, my=False):
    heading(doc, MY['treatment'] if my else 'Commercial Video Treatment', my)
    mood = ('Mood: humane confidence; vibe: precise but warm; emotional arc: disruption → recognition → agency → forward movement.' if not my else 'Mood: နွေးထွေးယုံကြည်မှု။ Vibe: တိကျသော်လည်း လူသားဆန်မှုရှိသည်။ Emotional arc: အနှောင့်အယှက် → သတိပြုမိခြင်း → ကိုယ်တိုင်ရွေးချယ်နိုင်မှု → ရှေ့ဆက်ခြင်း။')
    para(doc, mood, my)
    premise = c['video_premise'] if not my else 'Story Premise — ' + c['video_premise'] + ' Audience tension ကို ' + c['tension_my'] + ' မှ စတင်ပြီး message-recall ကို “' + c['tagline_my'] + '” ဖြင့် တည်ဆောက်ပါမည်။'
    para(doc, ('Story premise: ' + premise) if not my else premise, my)
    script = (
        f"VO: “{c['tagline']} Start with the work that matters. Ask the right question. Take the next step.” Dialogue is natural Myanmar/English as approved, never scripted as a testimonial. On-screen copy: “{c['message']}” / “{c['cta']}”."
    )
    if my:
        script = f"VO: “{c['tagline_my']} အရေးကြီးတဲ့အလုပ်ကနေ စတင်ပါ။ မှန်ကန်တဲ့ မေးခွန်းကို မေးပါ။ နောက်တစ်ဆင့်ကို ရွေးချယ်ပါ။” Dialogue ကို natural Myanmar/English အဖြစ် approval ရရှိမှသာ အသုံးပြုမည်ဖြစ်ပြီး testimonial အဖြစ် မရေးသားပါ။ On-screen copy: “{c['message_my']}” / “{c['cta_my']}”။"
    para(doc, script, my)
    shots = [
        ('01', '0–4s', 'A working moment meets an interruption cue.' if not my else 'လုပ်ငန်းလည်ပတ်နေချိန်တွင် interruption cue ပေါ်လာသည်။', 'Proposed/TBC ' + c['story_location'], 'Dawn / late afternoon', 'Macro → medium; 50mm; shallow DoF; slow push-in.', 'Talent pauses naturally; no distress performance. Practical work props; motivated window light; room tone.'),
        ('02', '4–9s', 'The audience tension becomes visible through one human decision.' if not my else 'Audience tension ကို လူတစ်ယောက်ရဲ့ ဆုံးဖြတ်ချက်တစ်ခုကနေ မြင်ရစေသည်။', 'Same / controlled set', 'Day', 'Medium profile; 35mm; soft DoF; lateral slider.', 'Natural glance to task; restrained blocking; neutral wardrobe; warm practicals; clean dialogue capture.'),
        ('03', '9–15s', 'A question reframes the problem.' if not my else 'မေးခွန်းတစ်ခုက problem ကို အသစ်ပြန်လည်ဖွင့်ဆိုသည်။', 'Proof or UI moment', 'Day', 'Over-shoulder; 50mm; focus rack; gentle tilt.', 'Hand marks critical task or taps UI. Hero prop is generic until client product clearance. Soft key and graphic accent.'),
        ('04', '15–22s', 'The proposed experience makes the next step tangible.' if not my else 'အဆိုပြု experience က နောက်တစ်ဆင့်ကို မြင်နိုင်စေသည်။', 'Demo / consultation zone', 'Day', 'Wide 24mm; deeper DoF; controlled orbit.', 'Facilitator guides, never hard-sells. Labelled proof surfaces, minimal furniture, accessible sightline.'),
        ('05', '22–30s', 'A calm proof detail earns attention.' if not my else 'တည်ငြိမ်တဲ့ proof detail က အာရုံစိုက်မှုကို ရယူသည်။', 'Technical detail', 'Day', 'Macro 85mm; shallow DoF; locked-off.', 'Hands only; no technical result implied. Practical LEDs; foley of click, paper, or interface tap.'),
        ('06', '30–38s', 'The protagonist chooses a client-approved next step.' if not my else 'Protagonist က client-approved နောက်တစ်ဆင့်ကို ရွေးချယ်သည်။', 'Booking / briefing moment', 'Day', 'Medium front; 50mm; gentle handheld stabilised move.', 'Consent is visible and unhurried. Clear bilingual copy; quiet score begins to resolve.'),
        ('07', '38–48s', 'The working day resumes with a feeling of agency—not a guarantee.' if not my else 'Guarantee မဟုတ်ဘဲ ကိုယ်တိုင်ရွေးချယ်နိုင်မှုကို ခံစားရစေသည့် လုပ်ငန်းနေ့ရက် ပြန်လည်ဆက်လက်သည်။', 'Original setting', 'Golden hour', 'Wide 35mm; moderate DoF; rising crane/gimbal.', 'Natural ensemble blocking, no staged customer claim. Warm grade, room ambience and restrained rhythmic edit.'),
        ('08', '48–60s', 'Tagline and CTA land with an accessible clear end card.' if not my else 'Accessible clear end card ဖြင့် Tagline နှင့် CTA ကို အဆုံးသတ်ပြသသည်။', 'Graphic end card', 'N/A', 'Static graphic; no lens; high-contrast hierarchy.', 'Client-approved logo/product only. Subtitles on; descriptive audio; 16:9/9:16/1:1 versions; legal line only after approval.'),
    ]
    table(doc, ['SHOT' if not my else 'Shot', 'DURATION' if not my else 'ကြာချိန်', 'STORY BEAT' if not my else 'Story Beat', 'LOCATION' if not my else 'Location', 'TIME' if not my else 'Time', 'CAMERA' if not my else 'Camera', 'DIRECTION / DESIGN / SOUND' if not my else 'Direction / Design / Sound'], shots, [0.35,0.5,1.2,0.95,0.55,1.2,2.0], my)
    heading(doc, 'Craft Direction' if not my else 'Craft Direction', my)
    craft = [
        ('Camera / framing', 'Cinema-style 4K capture; 24/35/50/85mm prime set; intentional shallow focus only where it preserves story clarity.'),
        ('Production design', 'Realistic Myanmar work texture, uncluttered surface language, generic/unbranded product placeholders until cleared, no imitation of competitor design.'),
        ('Wardrobe / talent', 'Everyday professional wardrobe; restrained palette; natural movement; no unsafe behaviour or unapproved technician portrayal.'),
        ('Lighting / colour', 'Warm functional practicals, controlled daylight, cobalt or solar-gold accent; natural skin tone; gentle contrast; avoid exaggerated greenwashing cues.'),
        ('Sound / music', 'Location ambience, clean boom/lav recording where dialogue is used, original/licensed minimalist pulse, foley, captions and optional audio description.'),
        ('Edit / VFX', 'Purposeful 2–4 second edit rhythm, match cuts and graphic callouts; only practical effects or approved subtle VFX; no fabricated dashboard or performance data.'),
        ('Accessibility / versions', 'Burned-in and sidecar Myanmar/English subtitles, contrast-tested supers, audio-description-ready master, 16:9 / 9:16 / 1:1 cutdowns, 6s / 15s / 30s / 60s deliverables.'),
        ('Masters / rights', '4K ProRes master, H.264 web masters, project archive, clean textless plate where feasible. Usage requires written approval for talent, music, location, product, logo, script and language territory/term.'),
    ]
    if my:
        craft = [
            ('Camera / Framing', 'Cinema-style 4K capture; 24/35/50/85mm prime set; story clarity ကိုမထိခိုက်စေသည့် intentional shallow focus ကိုသာ အသုံးပြုပါမည်။'),
            ('Production Design', 'Myanmar work texture အမှန်တကယ်ရှိစေပြီး client approval မတိုင်မီ generic/unbranded product placeholder ကိုသာ အသုံးပြုပါမည်။'),
            ('Wardrobe / Talent', 'နေ့စဉ် professional wardrobe; restrained palette; natural movement; unsafe behaviour သို့မဟုတ် unapproved technician portrayal မပြုလုပ်ပါ။'),
            ('Lighting / Colour', 'Warm functional practical၊ controlled daylight၊ cobalt သို့မဟုတ် solar-gold accent; natural skin tone; gentle contrast ကို အသုံးပြုပြီး exaggerated greenwashing cue ကိုရှောင်ပါမည်။'),
            ('Sound / Music', 'Location ambience၊ dialogue အတွက် clean boom/lav recording၊ original/licensed minimalist pulse၊ foley၊ caption နှင့် optional audio description ကို ထည့်သွင်းပါမည်။'),
            ('Edit / VFX', 'ရည်ရွယ်ချက်ရှိသော 2–4 second edit rhythm၊ match cut နှင့် graphic callout; approved subtle VFX သို့မဟုတ် practical effect ကိုသာ အသုံးပြုပြီး fabricated dashboard/performance data မပြုလုပ်ပါ။'),
            ('Accessibility / Version', 'Burned-in နှင့် sidecar Myanmar/English subtitle၊ contrast-tested super၊ audio-description-ready master၊ 16:9 / 9:16 / 1:1 cutdown၊ 6s / 15s / 30s / 60s deliverable များ ပါဝင်ပါမည်။'),
            ('Master / Rights', '4K ProRes master, H.264 web master, project archive နှင့် textless plate (ဖြစ်နိုင်လျှင်) ပါဝင်မည်။ Talent၊ music၊ location၊ product၊ logo၊ script နှင့် language territory/term ကို စာဖြင့် အတည်ပြုပြီးမှ အသုံးပြုမည်။'),
        ]
    table(doc, ['FIELD' if not my else 'အချက်', 'TREATMENT STANDARD' if not my else 'Treatment Standard'], craft, [1.45,5.25], my)
    doc.add_page_break()


def add_workflow_sources_ask(doc, c, my=False):
    heading(doc, MY['workflow'] if my else 'Workflow, Production Gates & Approval Matrix', my)
    gates = [
        ('01', 'Brief & strategy', 'ZYNTH strategy', 'Client approves tension, audience, objective, KPI and exclusions.'),
        ('02', 'Treatment', 'ZYNTH creative', 'Client approves concept territory, tagline, tone and proposed/TBCs.'),
        ('03', 'Script & storyboard', 'Production house / ZYNTH', 'Client legal/technical approves all claims, dialogue, supers and CTA.'),
        ('04', 'Casting / talent', 'Production house', 'Client approves talent, wardrobe, language suitability, release and term.'),
        ('05', 'Location & permit checks', 'Production house / client', 'Venue, authority and site owner confirmation; no assumption.'),
        ('06', 'Production design & engineering', 'Supplier / engineer', 'Build, electrical, safety, materials and load plan clearance.'),
        ('07', 'Call sheet', 'Production house', 'Client confirms key attendance, safety, schedule and escalation owner.'),
        ('08', 'Shoot / activation', 'Production house + ZYNTH PM', 'Live issue escalation and daily client sign-off.'),
        ('09', 'Data backup', 'Production house', 'Dual backup, checksum log and secure handoff.'),
        ('10', 'Offline / fine cut', 'Editor / ZYNTH', 'Creative and factual review; version control.'),
        ('11', 'Colour & sound', 'Post house', 'Client review of grade, mix, captions and accessibility.'),
        ('12', 'Legal / compliance review', 'Client legal / technical', 'Claims, privacy, rights and required disclosure sign-off.'),
        ('13', 'Language adaptation', 'Myanmar copy reviewer', 'Natural Unicode Myanmar client-facing copy approval.'),
        ('14', 'Master export / delivery', 'Post house', 'Approved master and cutdown inventory handoff.'),
        ('15', 'Archive & learnings', 'ZYNTH PM', 'Source archive, consent/rights log and KPI readout.'),
    ]
    if my:
        gates = [(a, {'Brief & strategy':'Brief & Strategy','Treatment':'Treatment','Script & storyboard':'Script & Storyboard','Casting / talent':'Casting / Talent','Location & permit checks':'Location & Permit Check','Production design & engineering':'Production Design & Engineering','Call sheet':'Call Sheet','Shoot / activation':'Shoot / Activation','Data backup':'Data Backup','Offline / fine cut':'Offline / Fine Cut','Colour & sound':'Colour & Sound','Legal / compliance review':'Legal / Compliance Review','Language adaptation':'Language Adaptation','Master export / delivery':'Master Export / Delivery','Archive & learnings':'Archive & Learnings'}[b], c, d) for a,b,c,d in gates]
    table(doc, ['GATE' if not my else 'Gate', 'DELIVERABLE' if not my else 'Deliverable', 'OWNER' if not my else 'Owner', 'APPROVAL / DEPENDENCY' if not my else 'Approval / Dependency'], gates, [0.4,1.65,1.25,3.25], my)
    heading(doc, MY['sources'] if my else 'Source Log & Sourced Assumptions', my)
    used = set(c['source_ids'])
    source_rows = [(sid, name, finding, url) for sid, name, finding, url in SOURCE_LOG if sid in used or sid == 'S05']
    table(doc, ['ID', 'SOURCE', 'RETAINED FINDING / LIMIT', 'URL'], source_rows, [0.4,1.3,3.0,1.85], my)
    assumption = 'Sourced context supports market relevance and Q4 timing guardrails. Audience sizes, conversion rates, incremental gross contribution, programme scope, package budgets, project schedule and break-even are ZYNTH planning assumptions only; they must be recalibrated to client data and supplier quotes.'
    if my:
        assumption = 'Sourced context သည် market relevance နှင့် Q4 timing guardrail ကိုသာ ထောက်ပံ့ပေးပါသည်။ Audience size၊ conversion rate၊ incremental gross contribution၊ programme scope၊ package budget၊ project schedule နှင့် break-even တို့သည် ZYNTH planning assumption များသာဖြစ်ပြီး client data နှင့် supplier quote များဖြင့် ပြန်လည်ညှိနှိုင်းရပါမည်။'
    para(doc, assumption, my, italic=True, color=COLORS['muted'], size=8.5)
    heading(doc, MY['ask'] if my else 'Approval Ask', my)
    ask = ('Approve or amend the commercial tension, target audience, conversion mechanism, Q4 timing window, Recommended planning envelope, measurement access, claims boundary and the proposed/TBC production route. Authorise the next gate only after client legal, technical, data, venue and procurement owners are named.')
    if my:
        ask = ('Commercial tension၊ target audience၊ conversion mechanism၊ Q4 timing window၊ Recommended planning envelope၊ measurement access၊ claims boundary နှင့် proposed/TBC production route ကို approve သို့မဟုတ် amend ပြုလုပ်ပေးပါ။ Client legal၊ technical၊ data၊ venue နှင့် procurement owner များကို အမည်တပ်ပြီးမှ နောက်တစ်ဆင့် gate ကို authorise ပြုလုပ်ပေးပါ။')
    para(doc, ask, my)
    table(doc, ['DECISION REQUESTED' if not my else 'လိုအပ်သော ဆုံးဖြတ်ချက်', 'CLIENT RESPONSE' if not my else 'Client Response'], [
        ('Proceed to discovery and supplier/venue verification' if not my else 'Discovery နှင့် supplier/venue verification သို့ ဆက်လက်သွားရန်', 'Approve / Amend / Hold' if not my else 'Approve / Amend / Hold'),
        ('Release only the agreed planning-envelope tranche' if not my else 'သဘောတူထားသည့် planning-envelope tranche ကိုသာ release ပြုလုပ်ရန်', 'Lean / Recommended / Flagship / TBC'),
        ('Confirm named approval owners and turnaround' if not my else 'Approval owner နှင့် turnaround ကိုအတည်ပြုရန်', 'Client to nominate'),
    ], [4.45, 2.25], my)
    para(doc, 'Unresolved TBCs: client brand/product, target lists, venue/site, dates, partner/creator/talent, production house, engineering, permits, rights, supplier quotes, data governance, claims evidence, schedule and signatory owners.' if not my else 'မဖြေရှင်းရသေးသော TBC များ — client brand/product၊ target list၊ venue/site၊ date၊ partner/creator/talent၊ production house၊ engineering၊ permit၊ rights၊ supplier quote၊ data governance၊ claims evidence၊ schedule နှင့် signatory owner များ။', my, italic=True, color=COLORS['red'], size=8.5)


def build_doc(c, my=False):
    doc = Document()
    configure_doc(doc)
    add_cover(doc, c, my)
    add_strategy(doc, c, my)
    add_plan(doc, c, my)
    add_commercials(doc, c, my)
    add_measurement(doc, c, my)
    add_design(doc, c, my)
    add_treatment(doc, c, my)
    add_workflow_sources_ask(doc, c, my)
    lang = 'MM-Hybrid' if my else 'EN'
    out = DOCS / lang / f"{c['id']}_{c['slug']}_{lang}.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    return out


def main():
    concepts = json.loads(DATA.read_text(encoding='utf-8'))
    created = []
    for c in concepts:
        created.append(build_doc(c, my=False))
        created.append(build_doc(c, my=True))
    manifest = ROOT / 'proposals' / 'proposal_generation_manifest.txt'
    manifest.write_text('\n'.join(str(p.relative_to(ROOT)) for p in created) + '\n', encoding='utf-8')
    print(f'Created {len(created)} Word documents.')
    for p in created:
        print(p)

if __name__ == '__main__':
    main()
