from __future__ import annotations

import json
from pathlib import Path
from docx import Document

ROOT=Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-energy-bilingual')
CONCEPTS=json.loads((ROOT/'data'/'concepts.json').read_text(encoding='utf-8'))
REPORT=ROOT/'validation'/'batch_validation.md'

EN_HEADINGS=['Executive Overview','Audience Insight','Audience Behaviour Change','Audience Journey','Seasonal / Special-Day Rationale','Experience / Campaign Plan','Content System','Talent Logic','Production Requirements','Three Planning Budget Packages','Scenario-Based Commercial Outcomes & Break-Even','KPI Scorecard & Monitoring','Risk, Compliance & Safety Register','Design / Experience Package','Commercial Video Treatment','Workflow, Production Gates & Approval Matrix','Source Log & Sourced Assumptions','Approval Ask']
MY_TOKENS=['အနှစ်ချုပ်','Audience Insight','လိုချင်သော','အချိန်အခါ','Content System','Production Requirements','MMK Budget Packages','KPI Scorecard','Risk / Compliance','Commercial Video Treatment','Approval Ask']

def doc_text(path):
    doc=Document(path)
    parts=[]
    for p in doc.paragraphs: parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells: parts.append(cell.text)
    return '\n'.join(parts)

def main():
    checks=[]
    errors=[]
    en=sorted((ROOT/'proposals'/'EN').glob('*.docx'))
    mm=sorted((ROOT/'proposals'/'MM-Hybrid').glob('*.docx'))
    checks.append(('English Word documents',len(en),10,len(en)==10))
    checks.append(('Hybrid Word documents',len(mm),10,len(mm)==10))
    checks.append(('Total standalone Word documents',len(en)+len(mm),20,len(en)+len(mm)==20))
    for path in en:
        tx=doc_text(path)
        missing=[h for h in EN_HEADINGS if h not in tx]
        ok=not missing and 'planning document' in tx.lower() and 'taxes' in tx.lower()
        checks.append((f'EN completeness: {path.name}','pass' if ok else 'fail','—' if ok else ', '.join(missing),ok))
        if not ok: errors.append(f'{path.name}: missing/invalid English sections: {missing}')
    for path in mm:
        tx=doc_text(path)
        my_chars=sum(1 for ch in tx if '\u1000'<=ch<='\u109f')
        missing=[t for t in MY_TOKENS if t not in tx]
        ok=my_chars>=500 and not missing
        checks.append((f'Hybrid language: {path.name}',f'{my_chars} Myanmar chars','—' if ok else ', '.join(missing),ok))
        if not ok: errors.append(f'{path.name}: insufficient Myanmar or missing labels: {missing}')
    unique_fields=['format','commercial_tension','conversion','territory','budget_structure','seasonal']
    for field in unique_fields:
        vals=[c[field] for c in CONCEPTS]
        ok=len(vals)==len(set(vals))==10
        checks.append((f'Concept uniqueness: {field}',len(set(vals)),10,ok))
        if not ok: errors.append(f'Duplicate {field}')
    ids=[c['id'] for c in CONCEPTS]
    checks.append(('Concept count',len(CONCEPTS),10,len(CONCEPTS)==10 and ids==[f'{i:02d}' for i in range(1,11)]))
    design_png=list((ROOT/'design_packages').rglob('*.png'));design_svg=list((ROOT/'design_packages').rglob('*.svg'))
    checks.append(('Client design previews',len(design_png),18,len(design_png)==18))
    checks.append(('Editable design SVG sources',len(design_svg),18,len(design_svg)==18))
    for file in ['monitoring/ZYNTH-20260821-ENERGY-Monitoring.xlsx','monitoring/monitoring_report.md','monitoring/source_manifest.json','research/verified_source_notes.md','validation/design_visual_qc.md']:
        p=ROOT/file;ok=p.exists() and p.stat().st_size>0;checks.append((f'Required file: {file}','present' if ok else 'missing','non-empty',ok));
        if not ok: errors.append(f'Missing required {file}')
    passed=sum(1 for *_,ok in checks if ok);total=len(checks)
    status='PASS' if passed==total else 'FAIL'
    out=['# Batch Validation Report','',f'> **Status:** {status} — {passed}/{total} checks passed.','', '| Check | Observed | Expected / detail | Status |','|---|---:|---|---|']
    for name,observed,expected,ok in checks:out.append(f"| {name} | {observed} | {expected} | {'PASS' if ok else 'FAIL'} |")
    out.extend(['','## Result','', 'The validation tests presence and structural completeness only. It does not convert planning assumptions into factual claims, supplier quotes, approvals, permits, product evidence, financial guarantees or legal clearance.'])
    if errors: out.extend(['','## Errors','']+[f'- {e}' for e in errors])
    REPORT.parent.mkdir(parents=True,exist_ok=True);REPORT.write_text('\n'.join(out)+'\n',encoding='utf-8')
    print(f'{status} {passed}/{total}')
    if errors: raise SystemExit(1)

if __name__=='__main__':main()
