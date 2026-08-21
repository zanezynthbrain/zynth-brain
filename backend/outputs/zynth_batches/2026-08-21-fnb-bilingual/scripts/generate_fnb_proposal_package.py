from pathlib import Path
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-fnb-bilingual')
CAMPAIGNS = json.loads((ROOT/'data/campaigns.json').read_text(encoding='utf-8'))['campaigns']
COMMERCIALS = json.loads((ROOT/'data/commercials.json').read_text(encoding='utf-8'))['commercials']
PROPOSALS = ROOT/'proposals'
STORYBOARDS = ROOT/'commercial_storyboards'
PROPOSALS.mkdir(exist_ok=True)
STORYBOARDS.mkdir(exist_ok=True)

NAVY = '153B4D'; TEAL = '1A6A68'; GOLD = 'C5973D'; PALE = 'EAF2F1'; INK = '1F2933'; GREY = '6B7280'

SOURCES = {
    'FNB-S01': ('USDA Foreign Agricultural Service — Burma: Food Service - Hotel Restaurant Institutional', 'https://www.fas.usda.gov/data/gain/2025/02/burma-food-service-hotel-restaurant-institutional'),
    'FNB-S02': ('Xinhua — Tastes of Golden Land', 'https://english.news.cn/20260726/49ff2f8a557b4825a58978e89ccccc61/c.html'),
    'FNB-S04': ('Myanmar Government — Upcoming Holidays', 'https://myanmar.gov.mm/upcoming-holidays'),
    'FNB-S05': ('Myanmar eVisa — Public Holidays in Myanmar 2026', 'https://evisa.moip.gov.mm/notice/public-holiday'),
    'FNB-S06': ('Tilleke & Gibbins — Myanmar-Language Labeling Required', 'https://www.tilleke.com/insights/myanmar-language-labeling-required-wide-range-products/'),
    'FNB-S07': ('Myanmar International TV — Tastes of the Golden Land', 'https://www.myanmaritv.com/news/tastes-golden-land-myanmar-traditional-food-competition-and-exhibition'),
    'FNB-S08': ('Ministry of Information — Traditional Food Competition and Exhibition 2026', 'https://www.moi.gov.mm/moi%3Aeng/news/21664'),
    'FNB-S09': ('Global New Light of Myanmar — Myanmar Urged to Cultivate Signature Foods', 'https://www.gnlm.com.mm/myanmar-urged-to-cultivate-signature-foods/'),
    'FNB-S10': ('Inter Myanmar Channel — Exploring HORECA 2026 Myanmar', 'https://www.youtube.com/watch?v=mbh1472c2HI'),
}

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=INK, size=8.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = 'Noto Sans Myanmar'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(.55); section.bottom_margin = Inches(.55)
    section.left_margin = Inches(.60); section.right_margin = Inches(.60)
    styles = doc.styles
    styles['Normal'].font.name = 'Noto Sans Myanmar'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    styles['Normal'].font.size = Pt(9)
    styles['Normal'].font.color.rgb = RGBColor.from_string(INK)
    for h in ['Title','Heading 1','Heading 2']:
        styles[h].font.name = 'Noto Sans Myanmar'
        styles[h]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar')
    styles['Title'].font.size = Pt(23); styles['Title'].font.color.rgb = RGBColor.from_string(NAVY)
    styles['Heading 1'].font.size = Pt(14); styles['Heading 1'].font.color.rgb = RGBColor.from_string(TEAL)
    styles['Heading 2'].font.size = Pt(10.5); styles['Heading 2'].font.color.rgb = RGBColor.from_string(NAVY)

def title_page(doc, title_en, title_mm, label, identifier):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run('ZYNTH')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(GOLD)
    p.add_run('  |  Myanmar-First Bilingual Production System').font.size = Pt(8.5)
    p = doc.add_paragraph(style='Title'); p.add_run(title_en)
    p = doc.add_paragraph(); r = p.add_run(title_mm); r.font.name = 'Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Noto Sans Myanmar'); r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(TEAL)
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f'{label}  •  {identifier}  •  F&B / Myanmar')
    r.bold = True; r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string(GREY)
    box = doc.add_table(rows=1, cols=1); box.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_cell_shading(box.cell(0,0), PALE)
    set_cell_text(box.cell(0,0), 'Client-ready planning proposal. All client, venue, chef, supplier, food-safety, product, claims, price, inventory, right, permit and partnership elements are proposed/TBC until written approval.', bold=True, color=NAVY, size=9)

def add_kv_table(doc, pairs):
    t = doc.add_table(rows=0, cols=2); t.alignment = WD_TABLE_ALIGNMENT.LEFT; t.style = 'Table Grid'
    for i,(k,v) in enumerate(pairs):
        row = t.add_row().cells
        set_cell_shading(row[0], PALE); set_cell_text(row[0], k, bold=True, color=NAVY)
        set_cell_text(row[1], v)
    doc.add_paragraph()

def add_table(doc, headers, rows, widths=None, font=8):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers):
        set_cell_shading(t.rows[0].cells[i], NAVY); set_cell_text(t.rows[0].cells[i], h, bold=True, color='FFFFFF', size=font)
    for n,rowdata in enumerate(rows):
        row=t.add_row().cells
        for i,val in enumerate(rowdata):
            if n % 2 == 0: set_cell_shading(row[i], 'F8FAFC')
            set_cell_text(row[i], val, size=font)
    doc.add_paragraph()
    return t

def heading(doc, english, myanmar):
    p = doc.add_paragraph(style='Heading 1'); p.add_run(english)
    p = doc.add_paragraph(); r=p.add_run(myanmar); r.italic=True; r.font.name='Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar'); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(GREY)

def add_bullets(doc, items):
    for item in items:
        p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.add_run(item)

def fmt_mmk(x): return f'MMK {x/1000000:.1f}m'

def campaign_doc(c):
    doc=Document(); style_doc(doc); title_page(doc,c['titleEn'],c['titleMm'],'CAMPAIGN / ACTIVATION PROPOSAL',c['id'])
    heading(doc,'1. Commercial Brief','စီးပွားရေးဆိုင်ရာ အကျဉ်းချုပ်')
    add_kv_table(doc,[('Format',c['format']),('Commercial tension',c['commercialTension']),('Objective',c['objective']),('Audience',c['audience']),('Creative territory',c['creativeTerritory']),('Seasonal / date logic',c['seasonalLogic'])])
    heading(doc,'2. Conversion Experience','ပြောင်းလဲမှု အတွေ့အကြုံ')
    add_kv_table(doc,[('Conversion mechanism',c['conversionMechanism']),('Experience flow',c['experienceFlow'])])
    heading(doc,'3. Proposed Workplan','အဆိုပြုလုပ်ငန်းစဉ်')
    add_table(doc,['Stage','Indicative timing','ZYNTH output','Client decision required'],[
        ('G0–G1: Brief & evidence','Working days 1–2','Evidence register, audience hypothesis, claims boundary','Approve objective, data boundary and product facts'),
        ('G2–G3: Strategy & creative','Working days 3–5','Journey, creative route, content/experience plan','Select route and priority conversion action'),
        ('G4: Feasibility','Working days 6–8','Food safety, hygiene, venue, rights, staffing, consent and risk register','Name accountable owners and clear go/no-go conditions'),
        ('G5: Production readiness','Working days 9–12','Run-of-show, asset list, training, capture and escalation plan','Approve final assets, rules, terms and operations'),
        ('G6–G7: Live & learn','Client/venue-confirmed','Dashboard check, daily log, decision record and learning note','Approve optimisation or close-out')
    ],font=8)
    heading(doc,'4. Budget Envelope','ဘတ်ဂျက် စီမံကိန်း')
    b=c['budgetMMK']; rec=b['recommended']; splits=[('Strategy, creative & project direction',.15),('Experience/production operations',.36),('Staffing, training & facilitation',.18),('Data, CRM & reporting',.08),('Product/sampling/guest operations',.13),('Contingency',.10)]
    add_table(doc,['Planning scenario','Non-tax envelope','Use'],[('Lean',fmt_mmk(b['lean']),'Minimum viable controlled pilot'),('Recommended',fmt_mmk(rec),b['structure']),('Flagship',fmt_mmk(b['flagship']),'Expanded scale, content and operational capacity')],font=8)
    add_table(doc,['Recommended budget workstream','Indicative share','Planning amount'],[(name,f'{share:.0%}',fmt_mmk(rec*share)) for name,share in splits],font=8)
    heading(doc,'5. Measurement and Commercial Scenarios','စောင့်ကြည့်တိုင်းတာမှုနှင့် Scenario')
    add_table(doc,['Scenario','Directional outcome'],[(k.title(),v) for k,v in c['planningOutcomes'].items()],font=8)
    add_kv_table(doc,[('Primary KPIs','; '.join(c['primaryKpis'])),('Data rule','Capture only client-approved, consented data needed for a stated follow-up; do not infer purchase intent, demographics or eligibility.'),('Scenario rule','Planning assumptions only. Replace with client funnel, margin, service capacity, inventory and attribution data before commercial sign-off.')])
    heading(doc,'6. Mandatory Preflight','မဖြစ်မနေ အတည်ပြုရမည့်အချက်များ')
    add_bullets(doc,c['requiredApprovals']+['No health, safety, origin, sustainability, quality, availability, price, award, restaurant, chef or partnership claim may be used without written client/legal approval.','All third-party visual, music, talent, recipe, venue and photography rights must be cleared before production.'])
    heading(doc,'7. Evidence and Use Limits','အထောက်အထားနှင့် အသုံးပြုခွင့် ကန့်သတ်ချက်')
    rows=[]
    for sid in c['sourceUse']:
        name,url=SOURCES[sid]; rows.append((sid,name,url))
    add_table(doc,['ID','Public source','URL'],rows,font=7.5)
    p=doc.add_paragraph(); p.add_run('Evidence use: ').bold=True; p.add_run('Sources provide public sector/event/calendar/compliance context only. They do not validate any future commercial partnership, venue, individual endorsement, performance, attendance or sales outcome.')
    doc.add_paragraph('Prepared by Manus AI for ZYNTH. Planning document only; not a contract, supplier quotation, legal advice or food-safety approval.',style='Caption')
    return doc

def commercial_doc(c):
    doc=Document(); style_doc(doc); title_page(doc,c['titleEn'],c['titleMm'],'STANDALONE COMMERCIAL + STORYBOARD DIRECTION',c['id'])
    heading(doc,'1. Film Brief','ဖիլմအကျဉ်းချုပ်')
    add_kv_table(doc,[('Format',c['format']),('Linked campaign',c['linkedCampaign']),('Commercial tension',c['tension']),('Objective',c['objective']),('Creative territory',c['territory']),('Storyboard status',c['storyboardStatus'])])
    heading(doc,'2. 12-Frame Detailed Storyboard','အသေးစိတ် Storyboard ၁၂ ခန်း')
    rows=[]
    for f in c['storyboard']:
        rows.append((f"{f['frame']:02d} | {f['timeWindow']}",f['visualAction'],f['camera'],f['audio'],f['onScreenText'],f['purpose']))
    add_table(doc,['Frame / time','Visual action','Camera','Audio','Text','Purpose'],rows,font=6.6)
    heading(doc,'3. Production & Claims Gate','ထုတ်လုပ်မှုနှင့် Claim အတည်ပြုချက်')
    add_bullets(doc,['This is a creative treatment and storyboard direction, not a final shoot script, product claim, talent confirmation or costed production bid.','Lock the final script only after client product facts, brand/legal review, food/allergen handling, location, talent, usage rights, music, subtitle language, pack shots, safety and end-card CTA are approved.','Prepare master, vertical, square and short cutdown needs only after the confirmed distribution and media plan is approved.','Use only consented real people; do not create customer testimonials, efficacy claims, venue endorsements or sponsor associations without proof and approval.'])
    heading(doc,'4. Production Planning Envelope','ထုတ်လုပ်မှု စီမံကိန်း')
    add_table(doc,['Workstream','Planning instruction'],[
        ('Development','Confirm factual proposition, audience, CTA, route and claim boundary.'),('Pre-production','Approve cast, location, food styling/handling, wardrobe, props, pack artwork, rights and shoot safety.'),('Production','Use a food-safe set plan, clear zones, client sign-off and a daily data/asset backup procedure.'),('Post-production','Build approved bilingual subtitles, versioning, music/licensing record, legal card and client review gates.'),('Delivery','Release only approved masters/cutdowns with an archive link and version status in the Master Tracker.')
    ],font=8)
    heading(doc,'5. Evidence Boundary','အထောက်အထား ကန့်သတ်ချက်')
    p=doc.add_paragraph('The film uses the F&B research register only as broad context. It does not represent any future event, location, association, product, person, brand, food claim, competition, sponsor or venue as confirmed. All such details remain proposed/TBC.')
    doc.add_paragraph('Prepared by Manus AI for ZYNTH. Creative planning document only; not a shooting permit, service contract, legal clearance or food-safety approval.',style='Caption')
    return doc

portfolio = ['# ZYNTH F&B Two-Hour Batch — Portfolio Overview','', '> **Batch:** `ZYNTH-20260821-FNB-BILINGUAL`  |  **Industry:** Food & beverage  |  **Status:** Research-grounded concepts; all client, commercial and production elements proposed/TBC.','', '## Ten Campaign / Activation Proposals','', '| ID | Campaign | Format | Conversion mechanism | Recommended planning envelope |','|---|---|---|---|---:|']
for c in CAMPAIGNS:
    portfolio.append(f"| {c['id']} | {c['shortName']} | {c['format']} | {c['conversionMechanism']} | {fmt_mmk(c['budgetMMK']['recommended'])} |")
portfolio += ['', '## Ten Separate Commercial / Storyboard Concepts','', '| ID | Commercial | Format | Linked campaign | Storyboard frames |','|---|---|---|---|---:|']
for c in COMMERCIALS:
    portfolio.append(f"| {c['id']} | {c['titleEn']} | {c['format']} | {c['linkedCampaign']} | {len(c['storyboard'])} |")
portfolio += ['', '## Research Sources', '']
for sid,(name,url) in SOURCES.items(): portfolio.append(f'- **{sid}:** [{name}]({url})')
portfolio += ['', '## Operating Note','', 'The campaign and commercial tracks are intentionally independent. Each campaign has a `CMP` ID; each commercial has a `COM` ID. Link them only where strategically useful, then keep budget, rights, production, approval and learning records separate in the live ZYNTH Master Tracker.']
(ROOT/'ZYNTH_FNB_Batch_Portfolio_Overview.md').write_text('\n'.join(portfolio)+'\n',encoding='utf-8')

for i,c in enumerate(CAMPAIGNS,1):
    campaign_doc(c).save(PROPOSALS/f'{i:02d}_{c["shortName"].lower().replace(" ","-").replace("&","and")}_Campaign_Proposal_MM-Bilingual.docx')
    md=['# '+c['titleEn'],'',c['titleMm'],'',f'**ID:** {c["id"]}  ',f'**Format:** {c["format"]}  ','', '## Commercial Brief','',f'**Tension:** {c["commercialTension"]}','',f'**Objective:** {c["objective"]}','',f'**Creative territory:** {c["creativeTerritory"]}','', '## Conversion Experience','',c['conversionMechanism'],'',c['experienceFlow'],'', '## Budget (MMK, non-tax planning envelope)','',f"Lean: {fmt_mmk(c['budgetMMK']['lean'])} | Recommended: {fmt_mmk(c['budgetMMK']['recommended'])} | Flagship: {fmt_mmk(c['budgetMMK']['flagship'])}",'', '## Primary KPIs','']+['- '+x for x in c['primaryKpis']]+['','## Required Approvals','']+['- '+x for x in c['requiredApprovals']]+['','> All client, venue, product, food-safety, claims, permits, rights and commercial details remain proposed/TBC. Scenario outcomes are planning assumptions, not guarantees.']
    (PROPOSALS/f'{i:02d}_{c["shortName"].lower().replace(" ","-").replace("&","and")}_Campaign_Proposal_MM-Bilingual.md').write_text('\n'.join(md)+'\n',encoding='utf-8')

for i,c in enumerate(COMMERCIALS,1):
    commercial_doc(c).save(STORYBOARDS/f'{i:02d}_{c["titleEn"].lower().replace(" ","-").replace(",","").replace("&","and")}_Commercial_Storyboard_MM-Bilingual.docx')
    md=['# '+c['titleEn'],'',c['titleMm'],'',f'**ID:** {c["id"]}  ',f'**Format:** {c["format"]}  ',f'**Linked campaign:** {c["linkedCampaign"]}  ','', '## Film Brief','',f'**Tension:** {c["tension"]}','',f'**Objective:** {c["objective"]}','',f'**Creative territory:** {c["territory"]}','', '## 12-Frame Storyboard','', '| # | Time | Visual action | Camera | Audio | Text | Purpose |','|---:|---|---|---|---|---|---|']
    for f in c['storyboard']:
        safe=[f['timeWindow'],f['visualAction'],f['camera'],f['audio'],f['onScreenText'],f['purpose']]
        md.append('| '+str(f['frame'])+' | '+' | '.join(x.replace('|','/') for x in safe)+' |')
    md += ['', '## Production Gate','', '> This is a creative storyboard direction only. Script, cast, locations, food/product facts, claims, pack shots, rights, music, subtitles, safety, end-card CTA and release plan require written client approval.']
    (STORYBOARDS/f'{i:02d}_{c["titleEn"].lower().replace(" ","-").replace(",","").replace("&","and")}_Commercial_Storyboard_MM-Bilingual.md').write_text('\n'.join(md)+'\n',encoding='utf-8')

print(f'Created {len(CAMPAIGNS)} campaign Word + Markdown proposals and {len(COMMERCIALS)} commercial Word + Markdown storyboards.')
