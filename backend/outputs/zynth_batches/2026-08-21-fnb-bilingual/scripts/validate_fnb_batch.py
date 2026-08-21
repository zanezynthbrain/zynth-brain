from pathlib import Path
import json
from docx import Document

ROOT=Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-fnb-bilingual')
campaigns=json.loads((ROOT/'data/campaigns.json').read_text(encoding='utf-8'))['campaigns']
commercials=json.loads((ROOT/'data/commercials.json').read_text(encoding='utf-8'))['commercials']
issues=[]; passed=[]

def check(cond,name,detail=''):
    (passed if cond else issues).append((name,detail))

check(len(campaigns)==10,'Campaign count',f'{len(campaigns)} records')
check(len(commercials)==10,'Commercial count',f'{len(commercials)} records')
check(sum(len(x['storyboard']) for x in commercials)==120,'Storyboard frame count',f'{sum(len(x["storyboard"]) for x in commercials)} frames')
for field in ['format','commercialTension','conversionMechanism','creativeTerritory','seasonalLogic']:
    vals=[x[field] for x in campaigns]
    check(len(set(vals))==10,f'Campaign unique {field}',f'{len(set(vals))}/10')
check(len(set(x['format'] for x in commercials))==10,'Commercial unique formats',f'{len(set(x["format"] for x in commercials))}/10')
check(len(set(x['territory'] for x in commercials))==10,'Commercial unique territories',f'{len(set(x["territory"] for x in commercials))}/10')

campaign_docs=sorted((ROOT/'proposals').glob('*.docx')); campaign_mds=sorted((ROOT/'proposals').glob('*.md'))
commercial_docs=sorted((ROOT/'commercial_storyboards').glob('*.docx')); commercial_mds=sorted((ROOT/'commercial_storyboards').glob('*.md'))
check(len(campaign_docs)==10,'Campaign Word document count',str(len(campaign_docs)))
check(len(campaign_mds)==10,'Campaign Markdown document count',str(len(campaign_mds)))
check(len(commercial_docs)==10,'Commercial Word document count',str(len(commercial_docs)))
check(len(commercial_mds)==10,'Commercial Markdown document count',str(len(commercial_mds)))

def doc_text(path):
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return '\n'.join(chunks)

for p in campaign_docs:
    t=doc_text(p)
    for token in ['Commercial Brief','Conversion Experience','Budget Envelope','Mandatory Preflight','Evidence and Use Limits']:
        check(token in t,f'Campaign heading {token}',p.name)
    check('proposed/TBC' in t,f'Campaign boundary disclaimer',p.name)
for p in commercial_docs:
    t=doc_text(p)
    for token in ['Film Brief','12-Frame Detailed Storyboard','Production & Claims Gate','Evidence Boundary']:
        check(token in t,f'Commercial heading {token}',p.name)
    check('creative treatment' in t.lower(),f'Commercial treatment disclaimer',p.name)

bad=['official partner','confirmed sponsor','award-winning','health benefit','safe for']
for p in campaign_mds+commercial_mds:
    text=p.read_text(encoding='utf-8').lower()
    for token in bad:
        check(token not in text,f'No unverified phrase: {token}',p.name)

lines=['# ZYNTH F&B Batch Validation','',f'**Batch:** `ZYNTH-20260821-FNB-BILINGUAL`  ',f'**Pass checks:** {len(passed)}  ',f'**Exceptions:** {len(issues)}  ','','## Result','', '**PASS**' if not issues else '**REVIEW REQUIRED**','', '## Checks Passed','']
for name,detail in passed: lines.append(f'- PASS — {name}: {detail}')
if issues:
    lines += ['', '## Exceptions', '']
    for name,detail in issues: lines.append(f'- REVIEW — {name}: {detail}')
(ROOT/'validation'/'batch_validation.md').write_text('\n'.join(lines)+'\n',encoding='utf-8')
print(f'PASS={len(passed)} EXCEPTIONS={len(issues)}')
if issues: raise SystemExit(1)
