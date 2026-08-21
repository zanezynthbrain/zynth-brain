from pathlib import Path
import json
from docx import Document
ROOT=Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-beauty-bilingual')
CAM=json.loads((ROOT/'data/campaigns.json').read_text(encoding='utf-8'))['campaigns']; COM=json.loads((ROOT/'data/commercials.json').read_text(encoding='utf-8'))['commercials']
checks=[]
def chk(name,condition,detail=''):
    checks.append((name,bool(condition),detail))
chk('Campaign count',len(CAM)==10,f'{len(CAM)}')
chk('Commercial count',len(COM)==10,f'{len(COM)}')
chk('Total storyboard frames',sum(len(c['storyboard']) for c in COM)==120,str(sum(len(c['storyboard']) for c in COM)))
chk('Unique campaign IDs',len({c['id'] for c in CAM})==10,'')
chk('Unique campaign formats',len({c['format'] for c in CAM})==10,'')
chk('Unique campaign territories',len({c['creativeTerritory'] for c in CAM})==10,'')
chk('Unique commercial IDs',len({c['id'] for c in COM})==10,'')
chk('Unique commercial titles',len({c['titleEn'] for c in COM})==10,'')
for c in CAM:
    chk(c['id']+' mandatory social/TikTok workstream',all(c['socialMediaPreparation'].get(k) for k in ['role','organicPreparation','tiktokPreparation','measurementPreparation','requiredSetup']),c['socialMediaPreparation'].get('role',''))
    chk(c['id']+' proposed/TBC boundary',('TBC' in c['audience'] or 'TBC' in c['seasonalLogic']) and len(c['requiredApprovals'])>=6,'')
    chk(c['id']+' source coverage',set(c['sourceUse'])=={'BEA-S01','BEA-S02','BEA-S03','BEA-S04'},'')
for c in COM:
    chk(c['id']+' exact 12 frames',len(c['storyboard'])==12,'')
    chk(c['id']+' linked campaign',c['linkedCampaign'] in {x['id'] for x in CAM},'')
    chk(c['id']+' storyboard field completeness',all(all(f.get(k) for k in ['frame','timeWindow','visualAction','camera','audio','onScreenText','purpose']) for f in c['storyboard']),'')
props=list((ROOT/'proposals').glob('*')); stories=list((ROOT/'commercial_storyboards').glob('*'))
chk('Campaign proposal file count',len(props)==20,str(len(props)))
chk('Commercial storyboard file count',len(stories)==20,str(len(stories)))
for p in sorted((ROOT/'proposals').glob('*.md')):
    t=p.read_text(encoding='utf-8'); chk(p.name+' social/TikTok heading','## Social Media & TikTok Preparation' in t,''); chk(p.name+' controlled wording','proposed/TBC' in t,'')
for p in sorted((ROOT/'commercial_storyboards').glob('*.md')):
    t=p.read_text(encoding='utf-8'); chk(p.name+' storyboard heading','## 12-Frame Storyboard' in t,''); chk(p.name+' controlled wording','written client approval' in t,'')
for p in sorted((ROOT/'proposals').glob('*.docx')):
    text='\n'.join(x.text for x in Document(p).paragraphs); chk(p.name+' Word social/TikTok section','Social Media & TikTok Preparation' in text,'')
for p in sorted((ROOT/'commercial_storyboards').glob('*.docx')):
    text='\n'.join(x.text for x in Document(p).paragraphs); chk(p.name+' Word storyboard section','12-Frame Detailed Storyboard' in text,'')
val=ROOT/'validation';val.mkdir(exist_ok=True); passed=sum(ok for _,ok,_ in checks); total=len(checks)
lines=['# ZYNTH Beauty / Cosmetics Batch Validation','',f'**Batch:** `ZYNTH-20260821-BEAUTY-BILINGUAL`  ','',f'**Result:** **{passed}/{total} PASS**  ','','| Check | Status | Detail |','|---|---|---|']
for name,ok,detail in checks:lines.append(f'| {name} | {"PASS" if ok else "FAIL"} | {detail} |')
lines+=['','## Scope Controls','','The 10 campaign/activation records explicitly contain a social-media and TikTok preparation workstream. Those workstreams do not replace the 10 independent commercial concepts or their 120 detailed storyboard frames. All product, claims, labels, creators, platform/account, retailer, safety, privacy, rights and commercial details remain proposed/TBC pending written client approval.']
(val/'batch_validation.md').write_text('\n'.join(lines)+'\n',encoding='utf-8')
print(f'{passed}/{total} PASS')
if passed!=total: raise SystemExit(1)
