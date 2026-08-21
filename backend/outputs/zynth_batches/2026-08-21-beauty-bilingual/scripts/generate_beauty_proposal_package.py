from pathlib import Path
import json, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/home/ubuntu/zynth-brain/backend/outputs/zynth_batches/2026-08-21-beauty-bilingual')
CAM=json.loads((ROOT/'data/campaigns.json').read_text(encoding='utf-8'))['campaigns']
COM=json.loads((ROOT/'data/commercials.json').read_text(encoding='utf-8'))['commercials']
PROP=ROOT/'proposals'; PROP.mkdir(exist_ok=True)
STORY=ROOT/'commercial_storyboards'; STORY.mkdir(exist_ok=True)
NAVY='2D2D2D'; PLUM='722F37'; ROSE='F4E5E8'; GOLD='B28A4A'; INK='1F2933'; GREY='6B7280'
SOURCES={'BEA-S01':('TikTok For Business — Beauty Advertising Guide','https://ads.tiktok.com/business/en-US/blog/beauty-advertising-tiktok-guide'),'BEA-S02':('Food and Drug Administration, Myanmar','https://www.fda.gov.mm/'),'BEA-S03':('TikTok Advertising Policies — Healthcare and Pharmaceuticals','https://ads.tiktok.com/resources/help/article/tiktok-ads-policy-healthcare-pharmaceuticals'),'BEA-S04':('Lebo Lion — TikTok Growth Tactics for Beauty Business Owners','https://www.youtube.com/watch?v=Wa7aHXpMw_4')}

def shade(cell,color):
    p=cell._tc.get_or_add_tcPr(); x=OxmlElement('w:shd'); x.set(qn('w:fill'),color); p.append(x)
def txt(cell,value,bold=False,color=INK,size=8.5):
    cell.text=''; p=cell.paragraphs[0]; p.paragraph_format.space_after=Pt(0); r=p.add_run(str(value)); r.bold=bold; r.font.name='Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar'); r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
def docstyle(d):
    s=d.sections[0]; s.top_margin=Inches(.55); s.bottom_margin=Inches(.55); s.left_margin=Inches(.58); s.right_margin=Inches(.58)
    for k in ['Normal','Title','Heading 1','Heading 2']:
        st=d.styles[k]; st.font.name='Noto Sans Myanmar'; st._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar')
    d.styles['Normal'].font.size=Pt(9); d.styles['Normal'].font.color.rgb=RGBColor.from_string(INK); d.styles['Title'].font.size=Pt(23); d.styles['Title'].font.color.rgb=RGBColor.from_string(NAVY); d.styles['Heading 1'].font.size=Pt(14); d.styles['Heading 1'].font.color.rgb=RGBColor.from_string(PLUM); d.styles['Heading 2'].font.size=Pt(10); d.styles['Heading 2'].font.color.rgb=RGBColor.from_string(NAVY)
def head(d,en,mm):
    d.add_paragraph(en,style='Heading 1'); p=d.add_paragraph(); r=p.add_run(mm); r.italic=True; r.font.name='Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar'); r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(GREY)
def title(d,en,mm,label,ident):
    p=d.add_paragraph(); r=p.add_run('ZYNTH'); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=RGBColor.from_string(GOLD); p.add_run('  |  Myanmar-First Bilingual Production System').font.size=Pt(8.5)
    d.add_paragraph(en,style='Title'); p=d.add_paragraph(); r=p.add_run(mm); r.font.name='Noto Sans Myanmar'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Noto Sans Myanmar'); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(PLUM)
    p=d.add_paragraph(); r=p.add_run(f'{label}  •  {ident}  •  Beauty / Cosmetics / Myanmar'); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor.from_string(GREY)
    t=d.add_table(rows=1,cols=1); shade(t.cell(0,0),ROSE); txt(t.cell(0,0),'Client-ready planning proposal. All brand, product, formulation, ingredient, label, claim, product result, before/after, retailer, marketplace, creator, talent, venue, social account, media budget, permit, rights, safety and partnership details are proposed/TBC until written approval.',True,NAVY,9)
def kv(d,pairs):
    t=d.add_table(rows=0,cols=2); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for a,b in pairs:
        row=t.add_row().cells; shade(row[0],ROSE); txt(row[0],a,True,NAVY); txt(row[1],b)
    d.add_paragraph()
def table(d,headers,rows,size=8):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.LEFT
    for i,h in enumerate(headers): shade(t.rows[0].cells[i],NAVY); txt(t.rows[0].cells[i],h,True,'FFFFFF',size)
    for n,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row):
            if n%2==0: shade(cells[i],'FAF8F9')
            txt(cells[i],v,False,INK,size)
    d.add_paragraph(); return t
def bullets(d,items):
    for x in items:
        p=d.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(2); p.add_run(x)
def money(x): return f'MMK {x/1e6:.1f}m'
def slug(s): return re.sub(r'[^a-z0-9]+','-',s.lower()).strip('-')

def campaign_doc(c):
    d=Document(); docstyle(d); title(d,c['titleEn'],c['titleMm'],'CAMPAIGN / ACTIVATION + SOCIAL PREPARATION PROPOSAL',c['id'])
    head(d,'1. Commercial Brief','စီးပွားရေးဆိုင်ရာ အကျဉ်းချုပ်'); kv(d,[('Format',c['format']),('Commercial tension',c['commercialTension']),('Objective',c['objective']),('Audience',c['audience']),('Creative territory',c['creativeTerritory']),('Seasonal/date logic',c['seasonalLogic'])])
    head(d,'2. Conversion Experience','ပြောင်းလဲမှု အတွေ့အကြုံ'); kv(d,[('Conversion mechanism',c['conversionMechanism']),('Experience flow',c['experienceFlow'])])
    head(d,'3. Social Media & TikTok Preparation','Social Media နှင့် TikTok အကြိုပြင်ဆင်မှု'); sm=c['socialMediaPreparation']; kv(d,[('Channel role',sm['role']),('Organic preparation',sm['organicPreparation']),('TikTok-specific preparation',sm['tiktokPreparation']),('Measurement preparation',sm['measurementPreparation'])]); table(d,['Preparation item','Required client confirmation'],[(x,'Confirm at G1/G4 before publishing, creator outreach, paid distribution or data capture.') for x in sm['requiredSetup']],7.5)
    head(d,'4. Proposed Workplan','အဆိုပြုလုပ်ငန်းစဉ်'); table(d,['Stage','Indicative timing','ZYNTH output','Client decision required'],[('G0–G1: Brief/evidence','Working days 1–2','Facts/claims boundary, audience hypothesis, social/TikTok guardrails','Approve factual product/claim, audience and privacy boundary'),('G2–G3: Strategy/creative','Working days 3–5','Activation journey, content pillars, TikTok hooks, CTA and moderation draft','Select route and controlled next action'),('G4: Feasibility','Working days 6–8','Hygiene, site, rights, creator, platform, data and risk register','Name accountable owners; clear go/no-go'),('G5: Production readiness','Working days 9–12','Run-of-show, content shot list, creator brief, asset/version plan','Approve final assets, terms and operations'),('G6–G7: Live/learn','Client-confirmed','Dashboard check, moderation log and decision record','Approve optimisation or close-out')],7.5)
    head(d,'5. Budget Envelope','ဘတ်ဂျက် စီမံကိန်း'); b=c['budgetMMK']; shares=[('Strategy, creative & project direction',.16),('Activation/production operations',.27),('Social/TikTok content preparation and versioning',.16),('Creator/rights/moderation preparation',.11),('Data, CRM & reporting',.10),('Client-approved product/guest operations',.10),('Contingency',.10)]; table(d,['Planning scenario','Non-tax envelope','Use'],[('Lean',money(b['lean']),'Minimum viable controlled pilot'),('Recommended',money(b['recommended']),b['structure']),('Flagship',money(b['flagship']),'Expanded approved content and operating capacity')],7.5); table(d,['Recommended workstream','Indicative share','Planning amount'],[(a,f'{q:.0%}',money(b['recommended']*q)) for a,q in shares],7.5)
    head(d,'6. Measurement and Commercial Scenarios','စောင့်ကြည့်တိုင်းတာမှုနှင့် Scenario'); table(d,['Scenario','Directional outcome'],[(a.title(),v) for a,v in c['planningOutcomes'].items()],7.5); kv(d,[('Primary KPIs','; '.join(c['primaryKpis'])),('Data rule','Capture only client-approved, consented data needed for a stated follow-up. Do not infer diagnosis, skin condition, demographics, purchase intent or eligibility.'),('Scenario rule','Planning assumptions only. Replace with client-validated funnel, margin, capacity, inventory, privacy and attribution data before sign-off.')])
    head(d,'7. Mandatory Beauty / Social Preflight','မဖြစ်မနေ အတည်ပြုရမည့်အချက်များ'); bullets(d,c['requiredApprovals']+['No efficacy, safety, medical, clinical, whitening, acne-treatment, transformation, before/after, testimonial, ingredient, dermatological, price, availability, retail, creator or platform-performance claim may be used without written approval.','No paid social/TikTok spend, platform targeting, creator posting, giveaway, product sample, UGC reuse or direct-message collection may occur before approved policy, rights, privacy, moderation and escalation controls.'])
    head(d,'8. Evidence and Use Limits','အထောက်အထားနှင့် အသုံးပြုခွင့် ကန့်သတ်ချက်'); table(d,['ID','Public source','URL'],[(s,SOURCES[s][0],SOURCES[s][1]) for s in c['sourceUse']],7); d.add_paragraph('Evidence use: public sources provide platform, agency-role and third-party workflow context only. They do not validate any client brand, formulation, product, claim, creator, retailer, account, sales, conversion, campaign outcome, regulator approval or platform entitlement.'); d.add_paragraph('Prepared by Manus AI for ZYNTH. Planning document only; not a contract, supplier quotation, legal/regulatory advice, product approval, platform approval or production clearance.',style='Caption'); return d

def commercial_doc(c):
    d=Document(); docstyle(d); title(d,c['titleEn'],c['titleMm'],'STANDALONE COMMERCIAL + STORYBOARD DIRECTION',c['id'])
    head(d,'1. Film Brief','ဖիլմအကျဉ်းချုပ်'); kv(d,[('Format',c['format']),('Linked campaign',c['linkedCampaign']),('Commercial tension',c['tension']),('Objective',c['objective']),('Creative territory',c['territory']),('Storyboard status',c['storyboardStatus'])])
    head(d,'2. 12-Frame Detailed Storyboard','အသေးစိတ် Storyboard ၁၂ ခန်း'); table(d,['Frame / time','Visual action','Camera','Audio','Text','Purpose'],[(f"{f['frame']:02d} | {f['timeWindow']}",f['visualAction'],f['camera'],f['audio'],f['onScreenText'],f['purpose']) for f in c['storyboard']],6.4)
    head(d,'3. Production, Social Versioning & Claims Gate','ထုတ်လုပ်မှု၊ Social Versioning နှင့် Claim အတည်ပြုချက်'); bullets(d,['This is a creative treatment and storyboard direction, not a final shoot script, product claim, confirmed cast/location, creator agreement or costed production bid.','Lock final script only after client product facts, brand/legal/regulatory review, pack artwork, talent, location, music, subtitles, social/TikTok format, rights, safety, privacy, approved end-card and CTA are cleared.','Prepare master, vertical 9:16, square, subtitles and short cutdowns only after confirmed distribution, platform eligibility and media plan approval.','Use fictional/non-product visual devices unless client approves product use, hygiene, sampling, safety, rights and claims. Do not create testimonials, treatment narratives, before/after comparisons, diagnosis or efficacy claims.'])
    head(d,'4. Production Planning Envelope','ထုတ်လုပ်မှု စီမံကိန်း'); table(d,['Workstream','Planning instruction'],[('Development','Confirm factual proposition, permitted claims, audience, CTA and content-release boundary.'),('Pre-production','Approve cast, creator role, location, wardrobe, props, pack artwork, product-handling plan, rights and shoot safety.'),('Production','Use a controlled set plan, hygiene/data/asset procedure, client sign-off and no unapproved product/medical narrative.'),('Post-production','Build approved bilingual subtitles, vertical versions, music/licensing record, legal card, claim log and review gates.'),('Delivery','Release approved masters/cutdowns only, with version status and content-rights record in the Master Tracker.')],7.5)
    head(d,'5. Evidence Boundary','အထောက်အထား ကန့်သတ်ချက်'); d.add_paragraph('The film uses the Beauty research register as broad platform, agency-role and workflow context. It does not represent a future brand, product, skin outcome, safety/efficacy result, creator, platform entitlement, retailer, venue, clinic, event or partner as confirmed. All such details remain proposed/TBC.'); d.add_paragraph('Prepared by Manus AI for ZYNTH. Creative planning document only; not a filming permit, product approval, platform approval, legal clearance or commercial contract.',style='Caption'); return d

portfolio=['# ZYNTH Beauty / Cosmetics Two-Hour Batch — Portfolio Overview','', '> **Batch:** `ZYNTH-20260821-BEAUTY-BILINGUAL`  |  **Industry:** Beauty / cosmetics  |  **Status:** Research-grounded concepts. Every campaign includes social-media and TikTok preparation; all brand, product and execution details remain proposed/TBC.','', '## Ten Integrated Campaign / Activation Proposals','', '| ID | Campaign | Format | Social/TikTok preparation | Recommended planning envelope |','|---|---|---|---|---:|']
for c in CAM: portfolio.append(f"| {c['id']} | {c['shortName']} | {c['format']} | {c['socialMediaPreparation']['role']} | {money(c['budgetMMK']['recommended'])} |")
portfolio+=['','## Ten Separate Commercial / Storyboard Concepts','','| ID | Commercial | Format | Linked campaign | Storyboard frames |','|---|---|---|---|---:|']
for c in COM: portfolio.append(f"| {c['id']} | {c['titleEn']} | {c['format']} | {c['linkedCampaign']} | {len(c['storyboard'])} |")
portfolio+=['','## Research Sources','']+[f'- **{s}:** [{n}]({u})' for s,(n,u) in SOURCES.items()]+['','## Operating Note','','Campaign and commercial tracks are intentionally independent. Social media and TikTok preparation is a mandatory workstream inside each campaign proposal; it does not replace the separate COM concept and 12-frame storyboard. No paid social/TikTok launch, creator outreach, product sampling, UGC reuse, data capture, platform targeting, clinic/procedure or product claim activity is authorised until the relevant client, legal/regulatory, platform, privacy, rights, safety and operations gates are approved.']
(ROOT/'ZYNTH_BEAUTY_Batch_Portfolio_Overview.md').write_text('\n'.join(portfolio)+'\n',encoding='utf-8')
for i,c in enumerate(CAM,1):
    base=f'{i:02d}_{slug(c["shortName"])}_Campaign_Proposal_MM-Bilingual'; campaign_doc(c).save(PROP/(base+'.docx'))
    sm=c['socialMediaPreparation']; md=['# '+c['titleEn'],'',c['titleMm'],'',f'**ID:** {c["id"]}  ',f'**Format:** {c["format"]}  ','','## Commercial Brief','',f'**Tension:** {c["commercialTension"]}','',f'**Objective:** {c["objective"]}','',f'**Creative territory:** {c["creativeTerritory"]}','', '## Conversion Experience','',c['conversionMechanism'],'',c['experienceFlow'],'','## Social Media & TikTok Preparation','',f'**Channel role:** {sm["role"]}','',f'**Organic preparation:** {sm["organicPreparation"]}','',f'**TikTok-specific preparation:** {sm["tiktokPreparation"]}','',f'**Measurement preparation:** {sm["measurementPreparation"]}','', '### Required Setup','']+['- '+x for x in sm['requiredSetup']]+['','## Budget (MMK, non-tax planning envelope)','',f'Lean: {money(c["budgetMMK"]["lean"])} | Recommended: {money(c["budgetMMK"]["recommended"])} | Flagship: {money(c["budgetMMK"]["flagship"])}','', '## Primary KPIs','']+['- '+x for x in c['primaryKpis']]+['','## Required Approvals','']+['- '+x for x in c['requiredApprovals']]+['','## Evidence and Use Limits','']+[f'- [{s}: {SOURCES[s][0]}]({SOURCES[s][1]})' for s in c['sourceUse']]+['','> All brand, product, formulation, claim, creator, venue, retailer, platform, rights, safety, privacy and commercial details remain proposed/TBC. TikTok/social preparation is a planning workstream; no post, paid distribution, data capture or creator outreach is authorised.']
    (PROP/(base+'.md')).write_text('\n'.join(md)+'\n',encoding='utf-8')
for i,c in enumerate(COM,1):
    base=f'{i:02d}_{slug(c["titleEn"])}_Commercial_Storyboard_MM-Bilingual'; commercial_doc(c).save(STORY/(base+'.docx'))
    md=['# '+c['titleEn'],'',c['titleMm'],'',f'**ID:** {c["id"]}  ',f'**Format:** {c["format"]}  ',f'**Linked campaign:** {c["linkedCampaign"]}  ','', '## Film Brief','',f'**Tension:** {c["tension"]}','',f'**Objective:** {c["objective"]}','',f'**Creative territory:** {c["territory"]}','', '## 12-Frame Storyboard','', '| # | Time | Visual action | Camera | Audio | Text | Purpose |','|---:|---|---|---|---|---|---|']
    md += ['| '+str(f['frame'])+' | '+' | '.join(str(f[x]).replace('|','/') for x in ['timeWindow','visualAction','camera','audio','onScreenText','purpose'])+' |' for f in c['storyboard']]
    md+=['','## Production, Social Versioning & Claims Gate','','> This is a creative storyboard direction only. Script, product facts, claims, pack shots, cast, location, music, subtitles, platform formats, safety, privacy, rights and end-card CTA require written client approval.']
    (STORY/(base+'.md')).write_text('\n'.join(md)+'\n',encoding='utf-8')
print(f'Created {len(CAM)} campaign Word + Markdown proposals and {len(COM)} commercial Word + Markdown storyboards.')
