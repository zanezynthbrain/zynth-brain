from pathlib import Path
from datetime import datetime, timezone
import json, math, textwrap
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from PIL import Image, ImageDraw, ImageFont

ROOT=Path('/home/ubuntu/zynth-brain')
OUT=ROOT/'daily_proposals'/'2026-08-21_Retail_Bilingual'
PROP=OUT/'proposals'; ASSET=OUT/'assets'; SOURCE=OUT/'sources'
for p in [PROP,ASSET,SOURCE]: p.mkdir(parents=True,exist_ok=True)
NOW='2026-08-21T00:12:00Z'

sources=[
 {'id':1,'title':'Myanmar Ministry of Foreign Affairs — Public Holidays','url':'https://www.mofa.gov.mm/about-myanmar/public-holidays/','date':'accessed 2026-08-21','supports':'Official 2026 public-holiday dates; Thadingyut 25–27 October 2026.'},
 {'id':2,'title':'6th Edition Myanmar Retail Sourcing Expo — 2026','url':'https://www.fibre2fashion.com/trade-fairs/6th-edition-myanmar-retail-sourcing-expo-2026-66372','date':'accessed 2026-08-21','supports':'Published 30 October–1 November 2026, Myanmar Plaza Yangon; retail/supplier/manufacturer audience and seminar context.'},
 {'id':3,'title':'Myanmar Economic Monitor, June 2026: Shock Amid Fragility','url':'https://reliefweb.int/report/myanmar/myanmar-economic-monitor-june-2026-shock-amid-fragility-enmy','date':'2026-06-16','supports':'Macro pressure, fuel shock, inflation, weak demand, operating uncertainty.'},
 {'id':4,'title':'Digital 2026: Myanmar','url':'https://datareportal.com/reports/digital-2026-myanmar','date':'2025-11-08','supports':'Planning context for internet, mobile, and social channels; late-2025 data used for 2026 planning.'},
 {'id':5,'title':'Myanmar Advertising Costs 2026','url':'https://marketingmyanmar.com/myanmar-advertising-costs-2026-facebook-google-tiktok-rate-card/','date':'2026-03-20','supports':'Indicative paid-media ranges; proxy-based and not a platform quotation.'},
 {'id':6,'title':'World Bank Myanmar Overview','url':'https://www.worldbank.org/en/country/myanmar/overview','date':'accessed 2026-08-21','supports':'Country context and first-party publication hub.'},
 {'id':7,'title':'ZYNTH internal planning knowledge — Myanmar event landscape','url':'https://github.com/zanezynthbrain/zynth-brain/blob/main/backend/knowledge/22_myanmar_event_landscape.md','date':'accessed 2026-08-21','supports':'Starting-point local production ranges; every item requires RFQ validation.'},
 {'id':8,'title':'ZYNTH internal planning knowledge — pricing and standards','url':'https://github.com/zanezynthbrain/zynth-brain/blob/main/backend/knowledge/09_pricing_and_standards.md','date':'accessed 2026-08-21','supports':'Agency fee, contingency, deposit, and reporting conventions; not supplier quotes.'},
]

concepts=[
 dict(n=1,slug='basket-reset',title='Basket Reset',my='ခြင်းတောင်းကို ပြန်စီ',form='Retail activation',territory='Relief from friction',tension='Price is the only comparison',mechanic='Retail scan',season='Year-end / New Year',behaviour='Build a planned basket and return within 30 days',cta='Scan your real basket. Leave with a better plan.',mycta='သင့်ရဲ့ တကယ်ဝယ်မယ့်ခြင်းတောင်းကို Scan လုပ်ပြီး ပိုကောင်းတဲ့အစီအစဉ်နဲ့ ပြန်ထွက်ပါ။',format='In-store navigation + scan-to-save activation',budget=(18000000,34000000,62000000),audience='Urban family shoppers and value-conscious young professionals',insight='When prices move quickly, shoppers compare line by line and lose confidence in the whole basket. A retailer can win by reducing decision effort, not by shouting the lowest price.',seasonal='A practical year-end reset; no public-holiday entitlement is assumed.',journey='See shelf marker → scan basket → receive three-option plan → checkout → 30-day return reminder.',msg='Make the whole basket feel understandable.',talent='Bilingual retail educators, not celebrity talent.',physical=True),
 dict(n=2,slug='the-curated-hour',title='The Curated Hour',my='ရွေးချယ်ပေးတဲ့ တစ်နာရီ',form='Pop-up',territory='Craft and mastery',tension='Choice is overwhelming',mechanic='Appointment booking',season='Festival gifting season',behaviour='Book a guided shopping appointment',cta='Book one calm hour for the gift that matters.',mycta='အရေးကြီးတဲ့ လက်ဆောင်အတွက် စိတ်အေးအေးနဲ့ တစ်နာရီကို ကြိုတင် Booking လုပ်ပါ။',format='Appointment-only gifting pop-up',budget=(22000000,42000000,78000000),audience='Gift buyers planning Thadingyut and year-end moments',insight='Gift shoppers do not need more options; they need confidence that the choice will be right for the recipient.',seasonal='Thadingyut is officially listed for 25–27 October 2026; appointment timing is proposed around that window, subject to retailer and venue confirmation.',journey='Short video → appointment slot → guided curation → gift-wrap → recipient follow-up.',msg='A better gift starts with less noise.',talent='Bilingual stylists and product specialists; no endorsement claim.',physical=True),
 dict(n=3,slug='source-to-shelf',title='Source to Shelf',my='အရင်းအမြစ်မှ စင်ပေါ်သို့',form='Exhibition',territory='Transparency',tension='Quality cannot be verified before buying',mechanic='Published performance record',season='No seasonal hook',behaviour='Choose a verified product and share its proof card',cta='See what made the shelf before you choose it.',mycta='မရွေးချယ်ခင် စင်ပေါ်ရောက်လာတဲ့ လမ်းကြောင်းကို ကြည့်ပါ။',format='Retail provenance exhibition + QR proof cards',budget=(26000000,52000000,96000000),audience='Quality-sensitive shoppers and supplier partners',insight='In uncertain conditions, provenance becomes a usable retail service: evidence lets customers spend with less regret.',seasonal='No special-day dependency; the exhibition can run as a modular always-on proof zone.',journey='Discover material/producer → scan proof → compare → purchase → submit review.',msg='Proof belongs on the shelf.',talent='Retail staff plus documented producer voices; consent required.',physical=True),
 dict(n=4,slug='retail-makers-forum',title='Retail Makers Forum',my='လက်လီလုပ်ငန်းရှင်များ ဖိုရမ်',form='Executive forum',territory='Competence and safety',tension='The service disappears after the sale',mechanic='Named-staff request',season='Harvest / year-end business cycle',behaviour='Bring a decision-maker and request a named follow-up',cta='Bring the person who owns the next decision.',mycta='နောက်ဆုံးဆုံးဖြတ်ချက်ကို တာဝန်ယူမယ့်သူကို အတူခေါ်လာပါ။',format='B2B retail forum + supplier matching',budget=(30000000,58000000,110000000),audience='Retail owners, procurement leads, landlords, and distributors',insight='Retail relationships weaken when nobody owns the next action. A forum can make accountability visible before the contract.',seasonal='Year-end planning logic is proposed; no venue or partner rights are assumed.',journey='ABM invitation → forum → named action owner → 14-day check-in → opportunity pipeline.',msg='The next move needs a name.',talent='Retail operators, invited case speakers, and bilingual moderator; all speakers TBC.',physical=True),
 dict(n=5,slug='try-before-trust',title='Try Before Trust',my='ယုံကြည်မီ စမ်းကြည့်',form='Experiential sampling',territory='Proof over promise',tension='The promise is not believed',mechanic='Trial unit',season='Monsoon',behaviour='Trial a product and opt into a reminder channel',cta='Try the difference. Keep the evidence.',mycta='ကွာခြားချက်ကို စမ်းကြည့်ပြီး သက်သေကို သိမ်းထားပါ။',format='Sampling bar + timed demo + Viber reminder opt-in',budget=(16000000,31000000,56000000),audience='Category switchers and hesitant first-time buyers',insight='A short, well-designed trial can replace a long claim when the product benefit is felt or seen.',seasonal='Monsoon is a planning cue only; exact timing depends on product category and indoor venue availability.',journey='See demo → sample → rate experience → opt into reminder → buy or request advisor.',msg='Trust can start in one small trial.',talent='Trained demonstrators and one product expert; no paid testimonial unless contracted.',physical=True),
 dict(n=6,slug='lights-on-together',title='Lights On Together',my='အတူတူ မီးထွန်းကြမယ်',form='Sponsorship activation',territory='Belonging',tension='The category is invisible',mechanic='Referral code',season='Thadingyut',behaviour='Refer a household and redeem a shared offer',cta='Bring one more home into the light.',mycta='အိမ်တစ်အိမ်ကို အလင်းထဲ ထပ်ခေါ်လာပါ။',format='Thadingyut community-lighting sponsorship activation',budget=(24000000,48000000,88000000),audience='Families and community-minded retail visitors',insight='A retailer can earn relevance when its offer helps people perform a meaningful seasonal act, not merely buy another item.',seasonal='The public holiday window is sourced; sponsorship property, municipal permissions, and partner rights are proposed/TBC.',journey='See community installation → scan referral → redeem family bundle → share light moment.',msg='The season feels brighter when the offer includes someone else.',talent='Community hosts and local creators, names TBC; consent and safeguarding required.',physical=True),
 dict(n=7,slug='learn-the-shelf',title='Learn the Shelf',my='စင်ပေါ်ကနေ သင်ယူ',form='Workshop series',territory='Curiosity and discovery',tension='Nobody explains anything',mechanic='Teach-back / comprehension check',season='School intake',behaviour='Attend a recurring session and complete a teach-back',cta='Come once. Leave able to explain it.',mycta='တစ်ကြိမ်လာပါ။ ပြန်ရှင်းပြနိုင်အောင် ထွက်သွားပါ။',format='Weekly micro-masterclass in-store',budget=(14000000,28000000,50000000),audience='New households, students, and practical learners',insight='Education is a conversion asset when it gives a customer language for a purchase they already want to make.',seasonal='School-intake logic is category-dependent and therefore proposed/TBC.',journey='Short explainer → register → attend → teach-back → product recommendation → repeat session.',msg='Understanding is part of the product.',talent='Product trainers and community educators; speaker roster TBC.',physical=True),
 dict(n=8,slug='return-the-value',title='Return the Value',my='တန်ဖိုးကို ပြန်ပေး',form='Community programme',territory='Generosity',tension='Convenience beat the relationship',mechanic='Community report-back',season='Year-end / New Year',behaviour='Return used items or packaging and make a repeat purchase',cta='Bring something back. See where the value goes.',mycta='တစ်ခုခု ပြန်ယူလာပါ။ တန်ဖိုးက ဘယ်ကိုရောက်သွားလဲ သိပါ။',format='Take-back points + monthly public report-back',budget=(20000000,39000000,72000000),audience='Repeat shoppers and sustainability-curious families',insight='A take-back promise only feels real when the retailer reports what happened after collection.',seasonal='Year-end is a proposed reset moment; collection partners, recycling claims, and permits are TBC.',journey='See collection point → return item → receive source-coded reward → view monthly report → return.',msg='Value is stronger when it comes back.',talent='Community facilitators; no environmental claim without verification.',physical=True),
 dict(n=9,slug='creator-cart',title='Creator Cart',my='ဖန်တီးသူနဲ့ ရွေးတဲ့ခြင်းတောင်း',form='Creator collaboration',territory='Recognition of the overlooked',tension='The first purchase never repeats',mechanic='Referral code',season='Festival gifting season',behaviour='Use a creator code and return for a second purchase',cta='A useful recommendation should survive the second visit.',mycta='အသုံးဝင်တဲ့ အကြံပြုချက်က ဒုတိယအကြိမ်လာတဲ့အထိ ခံနိုင်ရမယ်။',format='Creator-led shoppable content + store route',budget=(12000000,26000000,54000000),audience='Mobile-first shoppers who follow practical local creators',insight='Creators are most credible when they show the unglamorous second visit: what they bought again, not just what looked good once.',seasonal='Gifting season is a proposed content window; creator availability and rights are TBC.',journey='Creator video → saved cart → store route → referral redemption → 30-day re-purchase prompt.',msg='The best recommendation earns a second visit.',talent='Three to five local creators shortlisted after audience-quality audit; no names invented.',physical=False),
 dict(n=10,slug='one-more-reason',title='One More Reason',my='နောက်ထပ် အကြောင်းပြချက်တစ်ခု',form='Paid digital funnel',territory='Continuity and reliability',tension='The buyer fears being oversold',mechanic='Opt-in reminder channel',season='No seasonal hook',behaviour='Opt into a reminder and convert after a proof-led sequence',cta='No pressure. Just one useful reason at a time.',mycta='ဖိအားမပေးပါ။ တစ်ကြိမ်မှာ အသုံးဝင်တဲ့ အကြောင်းပြချက်တစ်ခုပဲ ပေးပါမယ်။',format='Paid digital funnel + CRM reactivation + content series',budget=(10000000,23000000,47000000),audience='Lapsed customers and high-intent browsers',insight='A careful follow-up sequence can make a retailer feel more reliable than a constant discount blast.',seasonal='No seasonal hook; always-on testing is recommended.',journey='Paid ad → landing page → proof content → opt-in → reminder → store or chat conversion.',msg='One useful reason at a time.',talent='No celebrity talent required; customer-service demonstrators only.',physical=False),
]

def money(n): return f'{n:,.0f} MMK'
def add_page_number(paragraph):
    run=paragraph.add_run()
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); run._r.append(fld)
def set_cell_shading(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)
def set_repeat_table_header(row):
    trPr=row._tr.get_or_add_trPr(); tblHeader=OxmlElement('w:tblHeader'); tblHeader.set(qn('w:val'),'true'); trPr.append(tblHeader)
def add_table(doc, headers, rows):
    t=doc.add_table(rows=1, cols=len(headers)); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.style='Table Grid'
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=str(h); set_cell_shading(c,'17365D');
        for p in c.paragraphs:
            for r in p.runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255); r.font.size=Pt(8)
    set_repeat_table_header(t.rows[0])
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cells[i].paragraphs:
                for r in p.runs: r.font.size=Pt(8)
    return t

def doc_setup(doc):
    sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.7); sec.right_margin=Inches(.7)
    styles=doc.styles
    styles['Normal'].font.name='Aptos'; styles['Normal'].font.size=Pt(9)
    for name,size,color in [('Title',28,'17365D'),('Heading 1',17,'17365D'),('Heading 2',12,'2F75B5')]:
        styles[name].font.name='Aptos'; styles[name].font.size=Pt(size); styles[name].font.bold=True; styles[name].font.color.rgb=RGBColor.from_string(color)
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; footer.add_run('ZYNTH | Retail proposal | '); add_page_number(footer)

def add_cover(doc,c,hybrid=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('ZYNTH').bold=True
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('CLIENT PITCH PROPOSAL'); r.font.size=Pt(12); r.font.bold=True; r.font.color.rgb=RGBColor(47,117,181)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(c['title']); r.font.size=Pt(26); r.font.bold=True; r.font.color.rgb=RGBColor(23,54,93)
    if hybrid:
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(c['my']); r.font.size=Pt(18); r.font.bold=True
        p=doc.add_paragraph('Myanmar-first hybrid edition | မြန်မာဘာသာကို ဦးစားပေးထားသော Hybrid Edition'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    else:
        p=doc.add_paragraph('Full English edition'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Retail / Shopping Malls | Myanmar | Planning document | 21 August 2026').italic=True
    doc.add_page_break()

def add_common(doc,c,hybrid=False):
    en=lambda x: x
    doc.add_heading('Executive decision',1)
    doc.add_paragraph(f'{c["title"]} is a {c["format"]} designed to move {c["audience"]} from {c["tension"].lower()} toward {c["behaviour"].lower()}. The concept is a planning recommendation, not a supplier quotation, financial guarantee, legal approval, permit confirmation, or confirmed partner right.')
    if hybrid: doc.add_paragraph('ဒီ Concept က ဝယ်ယူသူတွေမှာရှိတဲ့ အခက်အခဲကို လျှော့ချပေးပြီး လက်လီဆိုင်နဲ့ ပြန်လည်ဆက်သွယ်ဝယ်ယူစေဖို့ ရည်ရွယ်ထားတဲ့ Campaign ဖြစ်ပါတယ်။ ဈေးနှုန်းများသည် ခန့်မှန်းချက်သာဖြစ်ပြီး Supplier Quote မဟုတ်ပါ။')
    doc.add_heading('At a glance',1)
    add_table(doc,['Field','Recommendation'],[['Industry','Retail / Shopping Malls'],['Campaign / event type',c['format']],['Creative territory',c['territory']],['Commercial tension',c['tension']],['Conversion mechanism',c['mechanic']],['Seasonal logic',c['season']],['Primary behaviour change',c['behaviour']],['CTA',c['cta']]])
    doc.add_heading('Insight and behaviour change',1); doc.add_paragraph(c['insight'])
    if hybrid: doc.add_paragraph('ဝယ်သူတွေက ရွေးချယ်စရာများတာထက် မိမိအတွက် မှန်ကန်တဲ့ ဆုံးဖြတ်ချက်ကို ယုံကြည်စိတ်ချစွာ ချနိုင်ဖို့ လိုအပ်ပါတယ်။ ဒါကြောင့် ဒီ Campaign က ဝယ်ယူမှုကို လွယ်ကူရှင်းလင်းစေပြီး နောက်တစ်ကြိမ် ပြန်လာစေမယ့် အပြုအမူကို ဖန်တီးပါမယ်။')
    doc.add_heading('Audience journey and plan',1); doc.add_paragraph(c['journey'])
    if hybrid: doc.add_paragraph('Audience Journey — တွေ့မြင်ခြင်းမှ စတင်ပြီး Scan/Booking/Trial/Registration လုပ်ကာ ဝယ်ယူမှုနဲ့ ပြန်လည်ဆက်သွယ်မှုအထိ အဆင့်ဆင့် တည်ဆောက်ထားပါတယ်။')
    doc.add_heading('Message system and CTA examples',1)
    add_table(doc,['Layer','English','Myanmar-first hybrid'],[['Master line',''+c['msg'],c['my']],['CTA',c['cta'],c['mycta']],['Proof cue','Show the useful evidence before asking for the sale.','မေးခွန်းမတောင်းခင် အသုံးဝင်တဲ့ သက်သေကို အရင်ပြပါ။']])
    doc.add_heading('Seasonal rationale and sourced assumptions',1); doc.add_paragraph(c['seasonal'])
    doc.add_paragraph('Sourced assumptions: public facts are listed in the source log. Venue availability, sponsorship inventory, mall permissions, partner rights, talent availability, retail data, production pricing, and media performance are proposed/TBC until written confirmation. MMK budgets are planning ranges only.')
    doc.add_heading('Experience, campaign, and content system',1)
    add_table(doc,['Stage','Experience / content','Conversion evidence'],[['1. Attract','Hero key visual, short vertical video, retail media or paid social','Reach, 3-second views, landing-page sessions'],['2. Explain','Demo, proof card, workshop, appointment or creator route','Scan, registration, appointment, completion'],['3. Convert','Offer or service handoff with clear terms','Redeem, purchase, qualified lead'],['4. Continue','CRM reminder, report-back, re-purchase or referral','30/60-day return, opt-in, referral']])
    doc.add_paragraph('Video treatment: 30–45 second hero film opens on the shopper problem, shows the concept working in one continuous movement, then lands on the master line and a single CTA. Cutdowns: 15s, 6s, vertical stories, bilingual captions, and a silent retail-screen version. Usage rights for faces, music, locations, logos, product claims, and creator handles require written approval.')
    doc.add_heading('Talent logic and production workflow',1); doc.add_paragraph(c['talent'])
    add_table(doc,['Gate','Timing','Owner / approval'],[['Brief and claims lock','Week 1','Client + ZYNTH'],['Design, copy, route and risk review','Week 2','ZYNTH + client'],['Supplier RFQ, venue/site check, casting','Weeks 2–3','ZYNTH; client observes'],['Production and rehearsal / QA','Week 4','ZYNTH + suppliers'],['Launch / live execution','Week 5','Client go/no-go'],['Reporting and handover','Within 10 business days','ZYNTH + client CRM/sales']])
    doc.add_heading('Production requirements and design package',1)
    if c['physical']:
        doc.add_paragraph('Design format: labelled PNG package at 2400×1600 px, with hero perspective, front/stage elevation, plan/top view, and detail callouts where feasible. Source status: ZYNTH concept visual, not an engineering drawing. Approximate footprint: 6m × 4m unless site survey changes it. Materials: modular aluminium frame, reusable fabric graphics, painted MDF/plywood surfaces, vinyl floor graphics, LED or low-energy practical lighting, lockable storage, tablet/QR stands, queue stanchions, and bilingual wayfinding. Sightlines: keep a 1.2m clear circulation path; protect product and talent zones from queue conflict. Limitations: final dimensions, power, rigging, fire egress, loading access, mall house rules, accessibility, and structural safety require supplier engineering and venue approval. Lead time: 4–6 weeks; complexity: medium to high.')
        doc.add_paragraph('Furniture and zones: welcome/registration, product or proof zone, guided interaction zone, conversion/CRM desk, content capture point, and back-of-house storage. Approval asks: approve the design direction, site survey, claims list, materials sample, and RFQ release; do not treat this file as permit or supplier approval.')
    else:
        doc.add_paragraph('Digital-only experience storyboard: Frame 1 problem hook; Frame 2 proof or curation; Frame 3 product/service route; Frame 4 low-pressure CTA; Frame 5 CRM continuation. Deliverables: landing page wireframe, 9:16/1:1/4:5 cutdowns, bilingual captions, event or product feed, CRM sequence, measurement plan, and accessibility QA. Source status: UI storyboard, not a final build or media booking. Lead time: 3–5 weeks; complexity: medium.')
    doc.add_heading('Budget packages and commercial logic',1)
    lean,rec,flag=c['budget'];
    rows=[]
    for label,total,mix in [('Lean',lean,'Pilot, one location, limited content'),('Recommended',rec,'Integrated, measured, two-channel'),('Flagship',flag,'Modular scale, expanded content and reporting')]:
        direct=round(total/1.30); fee=round(direct*.15); cont=round(direct*.10); tax=total-direct-fee-cont
        rows.append([label,money(total),money(direct),money(fee),money(cont),money(tax),mix])
    add_table(doc,['Package','Total planning range','Pass-through direct costs','Agency fee (15% of direct)','Contingency (10% of direct)','Taxes / statutory allowance','Scope'],rows)
    doc.add_paragraph('Budget note: direct costs are pass-through planning allowances; agency fee is separately identified; taxes/statutory allowance is a placeholder pending legal/accounting treatment; contingency is not profit. Media spend, venue, fabrication, talent, travel, payment-gateway fees, permits, security, and third-party production are pass-throughs where applicable. Replace every line with written quotes before contract.')
    doc.add_heading('ROI scenarios and break-even logic',1)
    add_table(doc,['Scenario','Illustrative assumption','Commercial outcome'],[['Downside','60% of planned qualified actions; 70% of expected margin per action','Learning value retained; reallocate after first measurement gate'],['Base','100% of planned qualified actions; 100% of expected contribution margin','Break-even when incremental contribution margin equals total investment'],['Upside','140% of planned qualified actions; 110% of expected margin per action','Scale only after attribution and capacity checks']])
    doc.add_paragraph('Break-even formula: required incremental transactions or qualified opportunities = total campaign investment ÷ contribution margin per transaction or opportunity. The client must provide baseline traffic, average order value, gross margin, repeat rate, CRM reach, and sales close rate; no financial outcome is guaranteed in this planning document.')
    doc.add_heading('KPI scorecard',1)
    add_table(doc,['Objective','Leading KPI','Commercial KPI','Data owner'],[['Attention','Reach, views, footfall or registrations','Cost per qualified action','Media / event team'],['Consideration','Scan, completion, time, content saves','Qualified lead rate','ZYNTH + client'],['Conversion','Redemption, booking, trial-to-purchase','Incremental revenue or margin','Client sales / POS'],['Continuation','Opt-in, repeat, referral, report-back','30/60-day repeat contribution','Client CRM']])
    doc.add_heading('Risk, compliance, and cultural care',1)
    add_table(doc,['Risk','Level','Mitigation'],[['Unverified product or sustainability claim','High','Claims substantiation and client legal approval before publishing'],['Venue / mall permission or public-space issue','Medium/High','Written site approval; label all inventory proposed/TBC'],['Crowding, accessibility, fire egress','High','Site survey, queue plan, 1.2m clear path, supplier safety sign-off'],['Personal data and CRM consent','High','Explicit opt-in, purpose limitation, deletion and handover protocol'],['Talent / music / image rights','Medium','Written usage-rights schedule and release forms'],['Budget overrun','Medium','RFQ comparison, change control, contingency, weekly reconciliation']])
    doc.add_heading('Approval ask',1); doc.add_paragraph(f'Approve the strategic direction for {c["title"]}, the Recommended package of {money(rec)}, a discovery/site or data workshop, the claims-and-measurement checklist, and RFQ release. Confirm the client owner for brand, legal, retail operations, CRM, sales, and finance. Approve only the next gate; do not treat this as a quotation or permit.')
    if hybrid: doc.add_paragraph(f'Approval Ask — {c["title"]} ကို အခြေခံ Direction အဖြစ် အတည်ပြုပြီး Recommended package {money(rec)} အတွက် Discovery/Workshop၊ Claim စစ်ဆေးမှုနဲ့ RFQ ထုတ်ဖို့ ခွင့်ပြုပါ။ ဒီစာရွက်က Quotation သို့မဟုတ် Permit မဟုတ်ပါ။')
    doc.add_heading('Source log',1)
    add_table(doc,['ID','Source','Date','Use / limitation'],[[s['id'],s['title'],s['date'],s['supports']] for s in sources])
    doc.add_paragraph('References are public or internal planning sources. Public facts are not endorsements. Internal sources are planning references and do not replace supplier quotes, legal review, permit confirmation, or client data.')

def make_doc(c,hybrid):
    d=Document(); doc_setup(d); add_cover(d,c,hybrid); add_common(d,c,hybrid)
    edition = 'Hybrid_English_Myanmar' if hybrid else 'English'
    name=f"2026-08-21_ZYNTH_Retail_{c['n']:02d}_{c['slug']}_{edition}.docx"
    d.save(PROP/name); return str(PROP/name)

def font(size=20,bold=False):
    for p in ['/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf','/home/ubuntu/zynth-brain/backend/data/fonts/Inter-700.ttf']:
        if Path(p).exists(): return ImageFont.truetype(p,size=size)
    return ImageFont.load_default()
def make_asset(c):
    w,h=2400,1600; im=Image.new('RGB',(w,h),'#f4f1ea'); dr=ImageDraw.Draw(im); navy='#17365D'; blue='#2F75B5'; gold='#D6A84F'; dark='#202833'
    dr.rectangle((0,0,w,170),fill=navy); dr.text((80,45),f'ZYNTH RETAIL DESIGN PACKAGE | {c["title"]}',font=font(54,True),fill='white')
    if c['physical']:
        # hero perspective
        dr.polygon([(140,520),(900,360),(1600,500),(850,700)],fill='#dbe8f4',outline=navy)
        dr.polygon([(850,700),(1600,500),(1600,1050),(850,1250)],fill='#b8cee3',outline=navy)
        dr.polygon([(140,520),(850,700),(850,1250),(140,1080)],fill='#c7d8e9',outline=navy)
        dr.rectangle((460,590,1250,780),fill=gold,outline=dark,width=6); dr.text((520,650),'HERO / INTERACTION',font=font(34,True),fill=dark)
        dr.rectangle((930,420,1340,570),fill='#fffaf0',outline=dark,width=5); dr.text((960,465),'PROOF / CTA',font=font(28,True),fill=dark)
        dr.rectangle((230,880,520,1030),fill='#8aa3b8',outline=dark,width=5); dr.text((255,930),'WELCOME',font=font(25,True),fill='white')
        dr.text((80,1320),'Hero perspective | Front/elevation | Plan/top | Detail callouts',font=font(36,True),fill=navy)
        dr.text((80,1380),'Concept visual only — dimensions, structure, power, fire egress and supplier engineering TBC.',font=font(28),fill=dark)
        out=ASSET/f"2026-08-21_ZYNTH_Retail_{c['n']:02d}_{c['slug']}_Sketch_3D_Design_Package.png"
    else:
        frames=[('1. HOOK','Problem in one line'),('2. PROOF','Useful evidence'),('3. ROUTE','Choose / book / scan'),('4. CTA','Low-pressure next step'),('5. CRM','Return reason')]
        x0=110; y=430
        for i,(a,b) in enumerate(frames):
            x=x0+i*450; dr.rounded_rectangle((x,y,x+360,y+560),radius=25,fill='white',outline=blue,width=8); dr.rectangle((x,y,x+360,y+95),fill=blue); dr.text((x+25,y+28),a,font=font(25,True),fill='white'); dr.text((x+28,y+180),b,font=font(28,True),fill=dark)
        dr.text((80,1180),'Digital UI / experience storyboard | 9:16, 1:1, 4:5 cutdowns | bilingual captions',font=font(38,True),fill=navy)
        dr.text((80,1250),'Storyboard only — final UX, claims, tracking and media placements require approval.',font=font(28),fill=dark)
        out=ASSET/f"2026-08-21_ZYNTH_Retail_{c['n']:02d}_{c['slug']}_Digital_UI_Storyboard.png"
    im.save(out); return str(out)

# Generate docs/assets
paths=[]
for c in concepts:
    paths += [make_doc(c,False),make_doc(c,True)]
assets=[make_asset(c) for c in concepts]
# Source manifest
manifest={'batchCode':'ZYNTH-20260821-RETAIL-BILINGUAL','industryCode':'retail','selectedIndustry':'Retail / Shopping Malls','scheduledAt':NOW,'createdAt':NOW,'proposalCount':10,'documentCount':20,'videoConceptCount':10,'proposals':[{'n':c['n'],'title':c['title'],'myanmarTitle':c['my'],'format':c['format'],'budgetMMK':{'lean':c['budget'][0],'recommended':c['budget'][1],'flagship':c['budget'][2]},'documents':[str(PROP/f"2026-08-21_ZYNTH_Retail_{c['n']:02d}_{c['slug']}_English.docx"),str(PROP/f"2026-08-21_ZYNTH_Retail_{c['n']:02d}_{c['slug']}_Hybrid_English_Myanmar.docx")],'asset':assets[c['n']-1],'videoConcept':True,'sourceIds':[s['id'] for s in sources]} for c in concepts],'sources':sources,'notes':'All budget values are planning ranges, not supplier quotations or guarantees. Sponsorship, permits, partner rights, venues, talent, claims and inventory are proposed/TBC unless explicitly sourced.'}
(OUT/'source_manifest.json').write_text(json.dumps(manifest,ensure_ascii=False,indent=2),encoding='utf-8')
# Monitoring report
md=OUT/'2026-08-21_ZYNTH_Daily_Myanmar_Retail_Monitoring_Report.md'
md.write_text('# ZYNTH Daily Myanmar Retail Monitoring Report\n\n**Batch:** ZYNTH-20260821-RETAIL-BILINGUAL  \\n**Coverage:** 21 August 2026  \\n**Status:** Planning batch generated; live brand/social monitoring is limited to the public sources logged below.\n\n## Coverage\n\nThe batch covers official Myanmar public-holiday evidence, a dated Myanmar retail sourcing event listing, macroeconomic context, digital planning context, indicative Myanmar media pricing, country context, and ZYNTH planning standards. It does not claim private analytics, platform account access, verified brand contacts, confirmed venue availability, or supplier quotations.\n\n## Signals\n\nThe clearest dated seasonal signal is the official Thadingyut holiday window of 25–27 October 2026. The clearest retail-sector event signal is the published 30 October–1 November 2026 Myanmar Retail Sourcing Expo listing at Myanmar Plaza, Yangon. Both are used conservatively: seasonal timing is a rationale, while event participation, sponsorship, inventory and rights remain proposed/TBC.\n\n## Watchlist\n\nBefore client approval, verify retail footfall and POS baselines, category margin, CRM reach, mall/site availability, permit requirements, event organizer status, supplier RFQs, creator audience quality, product claims, and data-consent requirements.\n\n## Source log\n\n'+ '\n'.join([f'- [{s["id"]}] [{s["title"]}]({s["url"]}) — {s["date"]}; {s["supports"]}' for s in sources])+'\n',encoding='utf-8')
# Workbook
wb=Workbook(); ws=wb.active; ws.title='Proposal Index'; headers=['No.','Title','Myanmar title','Format','Lean MMK','Recommended MMK','Flagship MMK','Physical/Digital','Document count','Video concept','Asset','Status']
ws.append(headers)
for cell in ws[1]: cell.font=Font(bold=True,color='FFFFFF'); cell.fill=PatternFill('solid',fgColor='17365D'); cell.alignment=Alignment(wrap_text=True)
for c,a in zip(concepts,assets): ws.append([c['n'],c['title'],c['my'],c['format'],c['budget'][0],c['budget'][1],c['budget'][2],'Physical' if c['physical'] else 'Digital',2,'Yes',a,'Generated; validation pending'])
for col in ws.columns:
    letter=col[0].column_letter; ws.column_dimensions[letter].width=min(max(max(len(str(x.value or '')) for x in col)+2,12),36)
ws.freeze_panes='A2'; ws.auto_filter.ref=ws.dimensions
ws2=wb.create_sheet('Source Log'); ws2.append(['ID','Title','URL','Date','Supports']);
for s in sources: ws2.append([s['id'],s['title'],s['url'],s['date'],s['supports']])
ws3=wb.create_sheet('Budget Summary'); ws3.append(['Concept','Lean MMK','Recommended MMK','Flagship MMK','Pass-through / fee / tax / contingency','Break-even note'])
for c in concepts: ws3.append([c['title'],c['budget'][0],c['budget'][1],c['budget'][2],'Direct pass-through; 15% agency fee on direct; 10% contingency on direct; taxes/statutory allowance placeholder','Investment ÷ contribution margin per action'])
wb.save(OUT/'2026-08-21_ZYNTH_Retail_Monitoring_and_Proposal_Index.xlsx')
# claim registry
regp=ROOT/'backend/outputs/variety_registry.json'; reg=json.loads(regp.read_text(encoding='utf-8'))
reg['cycles'].append({'industry':'Retail & shopping malls','industryCode':'retail','run_id':NOW,'count':10,'batch_code':'ZYNTH-20260821-RETAIL-BILINGUAL'})
for c in concepts:
    reg['concepts'].append({'industry':'Retail & shopping malls','industryCode':'retail','n':c['n'],'form':c['form'].lower(),'territory':c['territory'].lower(),'tension':c['tension'].lower(),'behaviour':c['behaviour'].lower(),'mechanic':c['mechanic'].lower(),'budget_scale':['lean pilot','mid-weight integrated','flagship multi-market'][c['n']%3],'season':c['season'].lower(),'batch_code':'ZYNTH-20260821-RETAIL-BILINGUAL'})
reg['updated']=NOW; regp.write_text(json.dumps(reg,ensure_ascii=False,indent=1),encoding='utf-8')
print(json.dumps({'out':str(OUT),'documents':len(paths),'assets':len(assets),'manifest':str(OUT/'source_manifest.json'),'workbook':str(OUT/'2026-08-21_ZYNTH_Retail_Monitoring_and_Proposal_Index.xlsx')},ensure_ascii=False))
